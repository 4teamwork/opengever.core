from datetime import date
from datetime import datetime
from docx import Document
from docxcompose.properties import CustomProperties
from ftw.builder import Builder
from ftw.builder import create
from ftw.testbrowser import browsing
from ftw.testbrowser.pages import factoriesmenu
from ftw.testbrowser.pages import plone
from ftw.testbrowser.pages import statusmessages
from ftw.testbrowser.pages.statusmessages import error_messages
from ftw.testing import freeze
from opengever.document.docprops import TemporaryDocFile
from opengever.dossier.templatefolder import get_template_folder
from opengever.dossier.templatefolder.interfaces import ITemplateFolder
from opengever.journal.handlers import DOC_PROPERTIES_UPDATED
from opengever.journal.tests.utils import get_journal_entry
from opengever.kub.testing import KuBIntegrationTestCase
from opengever.ogds.base.actor import Actor
from opengever.testing import add_languages
from opengever.testing import FunctionalTestCase
from opengever.testing import IntegrationTestCase
from opengever.testing import solr_data_for
from opengever.testing import SolrIntegrationTestCase
from plone.portlets.constants import CONTEXT_CATEGORY
from plone.portlets.interfaces import ILocalPortletAssignmentManager
from plone.portlets.interfaces import IPortletAssignmentMapping
from plone.portlets.interfaces import IPortletManager
from unittest import skip
from zope.component import getMultiAdapter
from zope.component import getUtility
import json
import os


class TestDocumentWithTemplateFormPlain(IntegrationTestCase):

    @browsing
    def test_form_lists_all_templates_alphabetically_including_nested(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')

        expected_listing = [
            {'': '', 'Creator': 'nicole.kohler', 'Modified': '28.12.2010', 'Title': u'T\xc3\xb6mpl\xc3\xb6te Mit'},
            {'': '', 'Creator': 'nicole.kohler', 'Modified': '31.08.2016', 'Title': u'T\xc3\xb6mpl\xc3\xb6te Normal'},
            {'': '', 'Creator': 'nicole.kohler', 'Modified': '31.08.2016', 'Title': u'T\xc3\xb6mpl\xc3\xb6te Ohne'},
            {'': '', 'Creator': 'nicole.kohler', 'Modified': '29.02.2020', 'Title': u'T\xc3\xb6mpl\xc3\xb6te Sub'},
        ]

        self.assertEqual(expected_listing, browser.css('table.listing').first.dicts())

    @browsing
    def test_form_does_not_list_templates_from_dossiertemplates(self, browser):
        self.login(self.regular_user, browser)
        create(Builder('document')
               .titled(u'T\xc3\xb6mpl\xc3\xb6te in dossiertemplate')
               .with_dummy_content()
               .within(self.dossiertemplate))

        browser.open(self.dossier, view='document_with_template')
        expected_listing = [
            {'': '', 'Creator': 'nicole.kohler', 'Modified': '28.12.2010', 'Title': u'T\xc3\xb6mpl\xc3\xb6te Mit'},
            {'': '', 'Creator': 'nicole.kohler', 'Modified': '31.08.2016', 'Title': u'T\xc3\xb6mpl\xc3\xb6te Normal'},
            {'': '', 'Creator': 'nicole.kohler', 'Modified': '31.08.2016', 'Title': u'T\xc3\xb6mpl\xc3\xb6te Ohne'},
            {'': '', 'Creator': 'nicole.kohler', 'Modified': '29.02.2020', 'Title': u'T\xc3\xb6mpl\xc3\xb6te Sub'},
        ]

        self.assertEqual(expected_listing, browser.css('table.listing').first.dicts())

    @browsing
    def test_form_does_not_inlcude_participants_with_disabled_feature(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')

        # XXX - The CheckBoxWidget is rendered with two labels (only in testing).
        # The reason for that is a wrong widget renderer adapter lookup because
        # of the stacked globalregistry (see plone.testing.zca: pushGlobalRegistry
        # for more information).
        expected_labels = [
            u'Template',
            u'Filter',
            u'Title',
            u'Edit after creation',
        ]

        self.assertEqual(expected_labels, browser.css('#form label').text)

    @browsing
    def test_cancel_redirects_to_the_dossier(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')

        browser.find('Cancel').click()

        self.assertEqual(self.dossier, browser.context)
        self.assertEqual('tabbed_view', plone.view())

    @browsing
    def test_save_redirects_to_the_dossiers_document_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')

        browser.fill({
            'form.widgets.template': self.normal_template.UID(),
            'Title': 'Test Document',
            'Edit after creation': False,
        }).save()

        self.assertEqual(self.dossier, browser.context)
        self.assertEqual(self.dossier.absolute_url() + '#documents', browser.url)

    @browsing
    def test_new_document_is_titled_with_the_form_value(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')
        browser.fill({
            'form.widgets.template': self.normal_template.UID(),
            'Title': 'Test Document',
        }).save()

        document = self.dossier.listFolderContents()[-1]
        self.assertEqual('Test Document', document.title)

    @browsing
    def test_new_document_values_are_filled_with_default_values(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')
        browser.fill({
            'form.widgets.template': self.normal_template.UID(),
            'Title': 'Test Document',
        }).save()

        document = self.dossier.listFolderContents()[-1]
        self.assertEqual(date.today(), document.document_date)
        self.assertEqual(u'privacy_layer_no', document.privacy_layer)

    @browsing
    def test_file_of_the_new_document_is_a_copy_of_the_template(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')
        browser.fill({
            'form.widgets.template': self.normal_template.UID(),
            'Title': 'Test Document',
        }).save()

        document = self.dossier.listFolderContents()[-1]

        self.assertEqual(self.normal_template.file.data, document.file.data)
        self.assertNotEquals(self.normal_template.file, document.file)

    @browsing
    def test_doc_properties_are_not_created_when_disabled(self, browser):
        self.login(self.regular_user, browser)

        with freeze(datetime(2020, 10, 28, 0, 0)):
            browser.open(self.dossier, view='document_with_template')
            browser.fill({
                'form.widgets.template': self.asset_template.UID(),
                'Title': 'Test Docx',
            }).save()

        document = self.dossier.listFolderContents()[-1]
        self.assertEqual(u'Test Docx.docx', document.file.filename)

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual([], properties)

    @browsing
    def test_templates_without_a_file_are_not_listed(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')

        expected_documents = [
            u'T\xc3\xb6mpl\xc3\xb6te Mit',
            u'T\xc3\xb6mpl\xc3\xb6te Normal',
            u'T\xc3\xb6mpl\xc3\xb6te Ohne',
            u'T\xc3\xb6mpl\xc3\xb6te Sub',
        ]

        found_documents = [
            row.get('Title')
            for row in browser.css('table.listing').first.dicts()
        ]

        self.assertEqual(expected_documents, found_documents)
        self.assertNotIn(u'T\xc3\xb6mpl\xc3\xb6te Leer', found_documents)


class TestDocumentWithTemplateFormWithDocProperties(IntegrationTestCase):
    features = (
        'doc-properties',
    )

    maxDiff = None

    def assert_doc_properties_updated_journal_entry_generated(self, document, user):
        entry = get_journal_entry(document)

        self.assertEqual(DOC_PROPERTIES_UPDATED, entry['action']['type'])
        self.assertEqual(user.id, entry['actor'])
        self.assertEqual('', entry['comments'])

    @browsing
    def test_properties_are_added_when_created_from_template_with_doc_properties(self, browser):
        self.login(self.regular_user, browser)

        with freeze(datetime(2020, 9, 28, 0, 0)):
            browser.open(self.dossier, view='document_with_template')
            browser.fill({
                'form.widgets.template': self.docprops_template.UID(),
                'Title': 'Test Docx',
            }).save()

        document = self.dossier.listFolderContents()[-1]
        self.assertEqual(u'Test Docx.docx', document.file.filename)

        expected_doc_properties = {
            'Document.ReferenceNumber': 'Client1 1.1 / 1 / 44',
            'Document.SequenceNumber': '44',
            'Dossier.ReferenceNumber': 'Client1 1.1 / 1',
            'Dossier.Title': u'Vertr\xe4ge mit der kantonalen Finanzverwaltung',
            'User.FullName': u'B\xe4rfuss K\xe4thi',
            'User.ID': self.regular_user.id,
            'ogg.document.document_date': datetime(2020, 9, 28, 0, 0),
            'ogg.document.reference_number': 'Client1 1.1 / 1 / 44',
            'ogg.document.sequence_number': '44',
            'ogg.document.title': 'Test Docx',
            'ogg.document.version_number': 0,
            'ogg.dossier.external_reference': u'qpr-900-9001-\xf7',
            'ogg.dossier.reference_number': 'Client1 1.1 / 1',
            'ogg.dossier.sequence_number': '1',
            'ogg.dossier.title': u'Vertr\xe4ge mit der kantonalen Finanzverwaltung',
            'ogg.user.address1': 'Kappelenweg 13',
            'ogg.user.address2': 'Postfach 1234',
            'ogg.user.city': 'Vorkappelen',
            'ogg.user.country': 'Schweiz',
            'ogg.user.department': 'Staatskanzlei',
            'ogg.user.department_abbr': 'SK',
            'ogg.user.description': 'nix',
            'ogg.user.directorate': 'Staatsarchiv',
            'ogg.user.directorate_abbr': 'Arch',
            'ogg.user.email': 'foo@example.com',
            'ogg.user.email2': 'bar@example.com',
            'ogg.user.firstname': u'K\xe4thi',
            'ogg.user.lastname': u'B\xe4rfuss',
            'ogg.user.phone_fax': '012 34 56 77',
            'ogg.user.phone_mobile': '012 34 56 76',
            'ogg.user.phone_office': '012 34 56 78',
            'ogg.user.salutation': 'Frau',
            'ogg.user.title': u'B\xe4rfuss K\xe4thi',
            'ogg.user.url': 'http://www.example.com',
            'ogg.user.userid': self.regular_user.id,
            'ogg.user.zip_code': '1234',
            'ogg.document.creator.user.address1': 'Kappelenweg 13',
            'ogg.document.creator.user.address2': 'Postfach 1234',
            'ogg.document.creator.user.zip_code': '1234',
            'ogg.document.creator.user.salutation': 'Frau',
            'ogg.document.creator.user.directorate_abbr': 'Arch',
            'ogg.document.creator.user.phone_office': '012 34 56 78',
            'ogg.document.creator.user.email2': 'bar@example.com',
            'ogg.document.creator.user.department_abbr': 'SK',
            'ogg.document.creator.user.city': 'Vorkappelen',
            'ogg.document.creator.user.email': 'foo@example.com',
            'ogg.document.creator.user.userid': self.regular_user.id,
            'ogg.document.creator.user.phone_mobile': '012 34 56 76',
            'ogg.document.creator.user.phone_fax': '012 34 56 77',
            'ogg.document.creator.user.description': 'nix',
            'ogg.document.creator.user.url': 'http://www.example.com',
            'ogg.document.creator.user.title': u'B\xe4rfuss K\xe4thi',
            'ogg.document.creator.user.country': 'Schweiz',
            'ogg.document.creator.user.firstname': u'K\xe4thi',
            'ogg.document.creator.user.lastname': u'B\xe4rfuss',
            'ogg.document.creator.user.department': 'Staatskanzlei',
            'ogg.document.creator.user.directorate': 'Staatsarchiv',
        }

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(
                expected_doc_properties.items() + [('Test', 'Peter')],
                properties)
        self.assert_doc_properties_updated_journal_entry_generated(document, self.regular_user)

    @browsing
    def test_properties_are_added_when_created_from_template_without_doc_properties(self, browser):
        self.login(self.regular_user, browser)

        with freeze(datetime(2020, 10, 28, 0, 0)):
            browser.open(self.dossier, view='document_with_template')
            browser.fill({
                'form.widgets.template': self.asset_template.UID(),
                'Title': 'Test Docx',
            }).save()

        document = self.dossier.listFolderContents()[-1]
        self.assertEqual(u'Test Docx.docx', document.file.filename)

        expected_doc_properties = {
            'Document.ReferenceNumber': 'Client1 1.1 / 1 / 44',
            'Document.SequenceNumber': '44',
            'Dossier.ReferenceNumber': 'Client1 1.1 / 1',
            'Dossier.Title': u'Vertr\xe4ge mit der kantonalen Finanzverwaltung',
            'User.FullName': u'B\xe4rfuss K\xe4thi',
            'User.ID': self.regular_user.id,
            'ogg.document.document_date': datetime(2020, 10, 28, 0, 0),
            'ogg.document.reference_number': 'Client1 1.1 / 1 / 44',
            'ogg.document.sequence_number': '44',
            'ogg.document.title': 'Test Docx',
            'ogg.document.version_number': 0,
            'ogg.dossier.external_reference': u'qpr-900-9001-\xf7',
            'ogg.dossier.reference_number': 'Client1 1.1 / 1',
            'ogg.dossier.sequence_number': '1',
            'ogg.dossier.title': u'Vertr\xe4ge mit der kantonalen Finanzverwaltung',
            'ogg.user.address1': 'Kappelenweg 13',
            'ogg.user.address2': 'Postfach 1234',
            'ogg.user.city': 'Vorkappelen',
            'ogg.user.country': 'Schweiz',
            'ogg.user.department': 'Staatskanzlei',
            'ogg.user.department_abbr': 'SK',
            'ogg.user.description': 'nix',
            'ogg.user.directorate': 'Staatsarchiv',
            'ogg.user.directorate_abbr': 'Arch',
            'ogg.user.email': 'foo@example.com',
            'ogg.user.email2': 'bar@example.com',
            'ogg.user.firstname': u'K\xe4thi',
            'ogg.user.lastname': u'B\xe4rfuss',
            'ogg.user.phone_fax': '012 34 56 77',
            'ogg.user.phone_mobile': '012 34 56 76',
            'ogg.user.phone_office': '012 34 56 78',
            'ogg.user.salutation': 'Frau',
            'ogg.user.title': u'B\xe4rfuss K\xe4thi',
            'ogg.user.url': 'http://www.example.com',
            'ogg.user.userid': self.regular_user.id,
            'ogg.user.zip_code': '1234',
            'ogg.document.creator.user.address1': 'Kappelenweg 13',
            'ogg.document.creator.user.address2': 'Postfach 1234',
            'ogg.document.creator.user.zip_code': '1234',
            'ogg.document.creator.user.salutation': 'Frau',
            'ogg.document.creator.user.directorate_abbr': 'Arch',
            'ogg.document.creator.user.phone_office': '012 34 56 78',
            'ogg.document.creator.user.email2': 'bar@example.com',
            'ogg.document.creator.user.department_abbr': 'SK',
            'ogg.document.creator.user.city': 'Vorkappelen',
            'ogg.document.creator.user.email': 'foo@example.com',
            'ogg.document.creator.user.userid': self.regular_user.id,
            'ogg.document.creator.user.phone_mobile': '012 34 56 76',
            'ogg.document.creator.user.phone_fax': '012 34 56 77',
            'ogg.document.creator.user.description': 'nix',
            'ogg.document.creator.user.url': 'http://www.example.com',
            'ogg.document.creator.user.title': u'B\xe4rfuss K\xe4thi',
            'ogg.document.creator.user.country': 'Schweiz',
            'ogg.document.creator.user.firstname': u'K\xe4thi',
            'ogg.document.creator.user.lastname': u'B\xe4rfuss',
            'ogg.document.creator.user.department': 'Staatskanzlei',
            'ogg.document.creator.user.directorate': 'Staatsarchiv',
        }

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(
                expected_doc_properties.items(),
                properties)
        self.assert_doc_properties_updated_journal_entry_generated(document, self.regular_user)


class TestDocumentWithTemplateFormWithKuBContacts(KuBIntegrationTestCase):

    @browsing
    def test_error_message_is_displayed_when_kub_feature_active(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')
        self.assertEqual(
            ['The Contact and Authorities directory is only supported in the new UI.'],
            error_messages())

        with self.observe_children(self.dossier) as children:
            browser.fill({
                'form.widgets.template': self.normal_template.UID(),
                'Title': 'Test Docx',
            }).save()

        self.assertEqual(1, len(children['added']))
        doc = children['added'].pop()
        self.assertEqual(u'Test Docx', doc.title)


class TestDocumentWithTemplateFormWithOfficeConnector(IntegrationTestCase):
    features = (
        'officeconnector-checkout',
    )

    @browsing
    def test_opens_doc_with_officeconnector_when_feature_flaged(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.dossier, view='document_with_template')
        browser.fill({
            'form.widgets.template': self.normal_template.UID(),
            'Title': 'Test OfficeConnector',
        }).save()

        self.assertIn(
            "'oc:",
            browser.css('script.redirector').first.text,
            'OfficeConnector redirection script not found',
        )


class TestTemplateFolder(FunctionalTestCase):

    @browsing
    def test_adding(self, browser):
        self.grant('Manager')
        add_languages(['de-ch'])
        browser.login().open(self.portal)
        factoriesmenu.add('Template Folder')

        browser.fill({'Title (German)': u'Vorlagen',
                      'Title (English)': u'Templates'}).save()

        statusmessages.assert_no_error_messages()
        self.assertTrue(ITemplateFolder.providedBy(browser.context))

    @browsing
    def test_is_only_addable_by_manager(self, browser):
        browser.login().open(self.portal)

        self.grant('Administrator')
        browser.reload()
        self.assertNotIn(
            'Template Folder',
            factoriesmenu.addable_types()
        )

        self.grant('Manager')
        browser.reload()
        self.assertIn(
            'Template Folder',
            factoriesmenu.addable_types()
        )

    @browsing
    def test_manager_addable_types(self, browser):
        self.grant('Manager')
        templatefolder = create(Builder('templatefolder'))
        browser.login().open(templatefolder)

        self.assertEqual(
            ['Document', 'Task Template Folder', 'Template Folder'],
            factoriesmenu.addable_types())

    @skip("This test currently fails in a flaky way on CI."
          "See https://github.com/4teamwork/opengever.core/issues/3995")
    @browsing
    def test_supports_translated_title(self, browser):
        add_languages(['de-ch', 'fr-ch'])
        self.grant('Manager')
        browser.login().open()
        factoriesmenu.add('Template Folder')
        browser.fill({'Title (German)': u'Vorlagen',
                      'Title (French)': u'mod\xe8le'})
        browser.find('Save').click()

        browser.find(u'Fran\xe7ais').click()
        self.assertEqual(u'mod\xe8le', browser.css('h1').first.text)

        browser.find('Deutsch').click()
        self.assertEqual(u'Vorlagen', browser.css('h1').first.text)

    @browsing
    def test_do_not_show_dossier_templates_tab(self, browser):
        templatefolder = create(Builder('templatefolder'))

        browser.login().visit(templatefolder)

        self.assertEqual(0, len(browser.css('.formTab #tab-dossiertemplates')))

    @browsing
    def test_portlet_inheritance_is_blocked(self, browser):
        self.grant('Manager')
        add_languages(['de-ch'])
        browser.login().open(self.portal)
        factoriesmenu.add('Template Folder')
        browser.fill({'Title (German)': u'Vorlagen',
                      'Title (English)': u'Templates'}).save()

        statusmessages.assert_no_error_messages()
        self.assert_portlet_inheritance_blocked(
            'plone.leftcolumn', browser.context)

    @browsing
    def test_navigation_portlet_is_added(self, browser):
        self.grant('Manager')
        add_languages(['de-ch'])
        browser.login().open(self.portal)
        factoriesmenu.add('Template Folder')
        browser.fill({'Title (German)': u'Vorlagen',
                      'Title (English)': u'Templates'}).save()

        statusmessages.assert_no_error_messages()

        manager = getUtility(
            IPortletManager, name=u'plone.leftcolumn', context=browser.context)
        mapping = getMultiAdapter(
            (browser.context, manager), IPortletAssignmentMapping)

        portlet = mapping.get('navigation')

        self.assertIsNotNone(
            portlet, 'Navigation portlet not added to Templatefolder')
        self.assertEqual(0, portlet.topLevel)
        self.assertEqual('opengever-dossier-templatefolder', portlet.root)

    @browsing
    def test_portlets_are_inherited_on_sub_templatefolder(self, browser):
        templatefolder = create(Builder('templatefolder'))
        self.grant('Manager')
        add_languages(['de-ch'])
        browser.login().open(templatefolder)
        factoriesmenu.add('Template Folder')
        browser.fill({'Title (German)': u'Vorlagen',
                      'Title (English)': u'Templates'}).save()
        statusmessages.assert_no_error_messages()

        manager = getUtility(
            IPortletManager, name=u'plone.leftcolumn', context=browser.context)
        assignable = getMultiAdapter(
            (browser.context, manager), ILocalPortletAssignmentManager)
        self.assertFalse(assignable.getBlacklistStatus(CONTEXT_CATEGORY))


class TestTemplateFolderWithSolr(SolrIntegrationTestCase):

    @browsing
    def test_patch_template_folder_title_reindexes_containing_dossier(self, browser):
        self.login(self.administrator, browser)
        self.subtemplate.reindexObject()
        self.commit_solr()

        self.assertEqual(u'Templates new / Vorlagen neu',
                         solr_data_for(self.subtemplate)['containing_dossier'])

        data = {'title_de': 'Neuer Titel', 'title_en': 'New title'}
        browser.open(self.subtemplates.absolute_url(), data=json.dumps(data),
                     method='PATCH', headers=self.api_headers)

        self.assertEqual(204, browser.status_code)
        self.commit_solr()
        self.assertEqual(u'New title / Neuer Titel',
                         solr_data_for(self.subtemplate)['containing_dossier'])


class TestTemplateFolderMeetingEnabled(IntegrationTestCase):

    features = (
        'meeting',
    )

    @browsing
    def test_addable_types_with_meeting_feature(self, browser):
        self.login(self.manager, browser)
        browser.open(self.templates)

        expected_addable_types = [
            'Document',
            'Meeting Template',
            'Proposal Template',
            'Sablon Template',
            'Task Template Folder',
            'Template Folder',
        ]

        addable_types = factoriesmenu.addable_types()

        self.assertItemsEqual(expected_addable_types, addable_types)


class TestTemplateFolderUtility(FunctionalTestCase):

    def test_get_template_folder_returns_path_of_the_templatefolder(self):
        templatefolder = create(Builder('templatefolder'))

        self.assertEqual(templatefolder, get_template_folder())

    def test_get_template_folder_returns_allways_root_templatefolder(self):
        templatefolder = create(Builder('templatefolder'))
        create(Builder('templatefolder')
               .within(templatefolder))

        self.assertEqual(templatefolder, get_template_folder())


class TestTemplateFolderListings(SolrIntegrationTestCase):

    @browsing
    def test_renders_quickupload_uploadbox(self, browser):
        self.login(self.administrator, browser)  # admins can add templates
        browser.open(self.templates)

        self.assertIsNotNone(
            browser.css('#uploadbox').first,
            'Upload box should be available in template folders')

    @browsing
    def test_receipt_delivery_and_subdossier_column_are_hidden_in_document_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-documents')

        expected_table_header = [
            '',
            'Sequence number',
            'Title',
            'Author',
            'Document date',
            'Modification date',
            'Creation date',
            'Checked out by',
            'Disclosure status',
            'Reference number',
            'File extension',
            'Keywords',
        ]

        self.assertEqual(expected_table_header, browser.css('table.listing').first.lists()[0])

    @browsing
    def test_receipt_delivery_and_subdossier_column_are_hidden_in_sablon_template_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-sablontemplates')

        expected_table_header = [
            '',
            'Sequence number',
            'Title',
            'Author',
            'Document date',
            'Modification date',
            'Creation date',
            'Checked out by',
            'Disclosure status',
            'Reference number',
            'File extension',
            'Keywords',
        ]

        self.assertEqual(expected_table_header, browser.css('table.listing').first.lists()[0])

    @browsing
    def test_receipt_delivery_and_subdossier_column_are_hidden_in_proposal_templates_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-proposaltemplates')

        expected_table_header = [
            '',
            'Sequence number',
            'Title',
            'Author',
            'Document date',
            'Modification date',
            'Creation date',
            'Checked out by',
            'Disclosure status',
            'Reference number',
            'File extension',
            'Keywords',
        ]

        self.assertEqual(expected_table_header, browser.css('table.listing').first.lists()[0])

    @browsing
    def test_enabled_actions_are_limited_in_document_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-documents')

        expected_action_menu_content = [
            'Export as Zip',
            'Copy items',
            'Export selection',
            'Move items',
        ]

        self.assertItemsEqual(expected_action_menu_content, browser.css('.actionMenuContent li').text)

        self.login(self.administrator, browser)
        browser.open(self.templates, view='tabbedview_view-documents')

        expected_action_menu_content = [
            'Export as Zip',
            'Copy items',
            'Check in with comment',
            'Check in without comment',
            'Export selection',
            'Delete',
            'Move items',
        ]

        self.assertItemsEqual(expected_action_menu_content, browser.css('.actionMenuContent li').text)

        self.login(self.manager, browser)
        browser.open(self.templates, view='tabbedview_view-documents')

        self.assertItemsEqual(expected_action_menu_content, browser.css('.actionMenuContent li').text)

    @browsing
    def test_document_tab_lists_only_documents_directly_beneath(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-documents')

        templates = browser.css('table.listing').first.dicts(as_text=False)
        self.assertEqual(4, len(templates))

        expected_document_urls = [
            self.docprops_template.absolute_url(),
            self.asset_template.absolute_url(),
            self.normal_template.absolute_url(),
            self.empty_template.absolute_url(),
        ]

        document_urls = [
            element.get('Title').css('a').first.get('href')
            for element in templates
        ]

        self.assertEqual(expected_document_urls, document_urls)
        self.assertNotIn(self.subtemplate.absolute_url(), document_urls)

    @browsing
    def test_enabled_actions_are_limited_in_sablontemplates_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-sablontemplates')

        expected_actions = [
            'Export as Zip',
            'Copy items',
            'Export selection',
        ]

        self.assertItemsEqual(expected_actions, browser.css('.actionMenuContent li').text)

        self.login(self.administrator, browser)
        browser.open(self.templates, view='tabbedview_view-sablontemplates')

        expected_actions = [
            'Export as Zip',
            'Copy items',
            'Check in with comment',
            'Check in without comment',
            'Export selection',
            'Delete',
        ]

        self.assertItemsEqual(expected_actions, browser.css('.actionMenuContent li').text)

        self.login(self.manager, browser)
        browser.open(self.templates, view='tabbedview_view-sablontemplates')

        self.assertItemsEqual(expected_actions, browser.css('.actionMenuContent li').text)

    @browsing
    def test_enabled_actions_are_limited_in_proposaltemplates_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-proposaltemplates')

        expected_actions = [
            'Export as Zip',
            'Copy items',
            'Export selection',
        ]

        self.assertItemsEqual(expected_actions, browser.css('.actionMenuContent li').text)

        self.login(self.administrator, browser)
        browser.open(self.templates, view='tabbedview_view-proposaltemplates')

        expected_actions = [
            'Export as Zip',
            'Copy items',
            'Check in with comment',
            'Check in without comment',
            'Export selection',
            'Delete',
        ]

        self.assertItemsEqual(expected_actions, browser.css('.actionMenuContent li').text)

        self.login(self.manager, browser)
        browser.open(self.templates, view='tabbedview_view-proposaltemplates')

        self.assertItemsEqual(expected_actions, browser.css('.actionMenuContent li').text)


class TestTemplateFolderListingsSolr(SolrIntegrationTestCase):

    @browsing
    def test_sablontemplates_tab_lists_only_documents_directly_beneath(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-sablontemplates')

        templates = browser.css('table.listing').first.dicts(as_text=False)
        self.assertEqual(1, len(templates))

        document_link = templates[0]['Title'].css('a').first.get('href')
        self.assertEqual(self.sablon_template.absolute_url(), document_link)

    @browsing
    def test_proposaltemplates_tab_lists_only_documents_directly_beneath(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-proposaltemplates')

        templates = browser.css('table.listing').first.dicts(as_text=False)
        self.assertEqual(3, len(templates))

        document_link = templates[0]['Title'].css('a').first.get('href')
        self.assertEqual(self.proposal_template.absolute_url(), document_link)


class TestTemplateDocumentTabs(IntegrationTestCase):

    features = ('bumblebee', )

    @browsing
    def test_visible_tabs(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbed_view')
        expected_tabs = ['Documents', 'Task template folders']
        self.assertEqual(expected_tabs, browser.css('.tabbedview-tabs span').text)

    @browsing
    def test_template_overview_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.normal_template, view='tabbedview_view-overview')

        self.assertIn(['Title', u'T\xc3\xb6mpl\xc3\xb6te Normal'], browser.css('table.listing').first.lists())

    @browsing
    def test_template_journal_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.normal_template, view='tabbedview_view-journal')

        journal_entries = browser.css('table.listing').first.dicts()

        self.assertEqual(Actor.lookup(self.administrator.id).get_label(), journal_entries[0]['Changed by'])
        self.assertEqual(u'Document added: T\xc3\xb6mpl\xc3\xb6te Normal', journal_entries[0]['Title'])


class TestTemplateDocumentTabsSolr(SolrIntegrationTestCase):

    features = ('bumblebee', )

    @browsing
    def test_documents_tab_shows_only_docs_directly_inside_the_folder(self, browser):
        self.login(self.regular_user, browser=browser)

        # main templatefolder
        browser.open(self.templates, view='tabbedview_view-documents-gallery')
        self.assertEqual(
            [self.empty_template.title, self.normal_template.title,
             self.asset_template.title, self.docprops_template.title],
            [img.attrib.get('alt') for img in browser.css('img')])

        # subtemplatefolder
        browser.open(self.subtemplates, view='tabbedview_view-documents-gallery')
        self.assertEqual(
            [self.subtemplate.title],
            [img.attrib.get('alt') for img in browser.css('img')])


class TestTemplateDocumentTabsWithOneoffixx(IntegrationTestCase):

    features = ('oneoffixx', )

    @browsing
    def test_visible_tabs(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbed_view')
        expected_tabs = ['OneOffixx', 'Documents', 'Task template folders']
        self.assertEqual(expected_tabs, browser.css('.tabbedview-tabs span').text)


class TestDossierTemplateFeature(IntegrationTestCase):
    features = (
        'dossiertemplate',
    )

    @browsing
    def test_visible_tabs(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbed_view')
        expected_tabs = ['Documents', 'Dossier templates', 'Task template folders']
        self.assertEqual(expected_tabs, browser.css('.tabbedview-tabs span').text)

    @browsing
    def test_dossier_template_is_addable_if_dossier_template_feature_is_enabled(self, browser):
        self.login(self.manager, browser)
        browser.open(self.templates)

        expected_addable_types = ['Document',
                                  'Dossier template',
                                  'Task Template Folder',
                                  'Template Folder']
        self.assertEqual(expected_addable_types, factoriesmenu.addable_types())

    @browsing
    def test_show_dossier_templates_tab(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates)

        self.assertEqual('Dossier templates', browser.css('.formTab #tab-dossiertemplates').first.text)


class TestTaskTemplateFoldersTab(SolrIntegrationTestCase):

    @browsing
    def test_tasktemplatefolders_tab_does_not_show_subtasktemplatefolders(self, browser):
        self.activate_feature('tasktemplatefolder_nesting')

        self.login(self.administrator, browser=browser)
        browser.open(self.tasktemplatefolder)
        factoriesmenu.add(u'Task Template Folder')
        browser.fill({'Title': 'Baugesuch', 'Type': 'parallel'}).submit()
        self.commit_solr()

        self.login(self.regular_user, browser=browser)
        browser.open(self.templates, view='tabbedview_view-tasktemplatefolders')
        templates = browser.css('table.listing').first.dicts()
        self.assertEqual(1, len(templates))
        self.assertEqual(templates[0]["Title"], self.tasktemplatefolder.Title())


class TestTemplateFolderShowroomPreviews(SolrIntegrationTestCase):

    features = ('meeting', 'bumblebee',)

    @staticmethod
    def get_expected_urls_from_documents(document_list):
        return [os.path.join(document.absolute_url(), "@@bumblebee-overlay-listing")
                for document in document_list]

    @browsing
    def test_document_tab_contains_showroom_preview_links(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-documents')
        expected_documents = [self.docprops_template, self.asset_template,
                              self.normal_template, self.empty_template]

        expected_urls = self.get_expected_urls_from_documents(expected_documents)
        actual_urls = [element.get("data-showroom-target") for element in browser.css(".showroom-item")]
        self.assertEqual(expected_urls, actual_urls)

    @browsing
    def test_sablontemplates_tab_contains_showroom_preview_links(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-sablontemplates')
        expected_documents = [self.sablon_template]

        expected_urls = self.get_expected_urls_from_documents(expected_documents)
        actual_urls = [element.get("data-showroom-target") for element in browser.css(".showroom-item")]
        self.assertEqual(expected_urls, actual_urls)

    @browsing
    def test_proposaltemplates_tab_contains_showroom_preview_links(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.templates, view='tabbedview_view-proposaltemplates')
        expected_documents = [self.proposal_template,
                              self.ad_hoc_agenda_item_template,
                              self.recurring_agenda_item_template]

        expected_urls = self.get_expected_urls_from_documents(expected_documents)
        actual_urls = [element.get("data-showroom-target") for element in browser.css(".showroom-item")]
        self.assertEqual(expected_urls, actual_urls)
