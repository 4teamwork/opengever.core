from docx import Document
from docxcompose.properties import CustomProperties
from ftw.builder import Builder
from ftw.builder import create
from opengever.document.docprops import TemporaryDocFile
from opengever.document.versioner import Versioner
from opengever.dossier.behaviors.dossier import IDossier
from opengever.dossier.behaviors.dossier import IDossierMarker
from opengever.dossier.command import CreateDocumentFromTemplateCommand
from opengever.dossier.command import CreateDossierFromTemplateCommand
from opengever.testing import IntegrationTestCase


class TestCreateDocumentFromTemplateCommand(IntegrationTestCase):

    features = ('doc-properties',)

    def test_create_document_from_template(self):
        expected_title = 'My title'
        expected_data = 'Test data'
        self.login(self.regular_user)
        template = create(Builder('document').within(self.dossiertemplate).with_dummy_content())
        command = CreateDocumentFromTemplateCommand(self.dossier, template, expected_title)
        document = command.execute()
        self.assertEqual(expected_title, document.title)
        self.assertEqual(expected_data, document.file.data)

    def test_create_document_from_template_without_file(self):
        expected_title = 'My title'
        self.login(self.regular_user)
        template = create(Builder('document').within(self.dossiertemplate))
        command = CreateDocumentFromTemplateCommand(self.dossier, template, expected_title)
        document = command.execute()
        self.assertEqual(expected_title, document.title)
        self.assertIsNone(document.file)

    def test_create_document_from_template_with_keywords(self):
        expected_title = 'My title'
        expected_keywords = ('My keyword1', 'My keyword2')
        self.login(self.regular_user)
        template = create(Builder('document'))
        command = CreateDocumentFromTemplateCommand(self.dossier, template, expected_title, keywords=expected_keywords)
        document = command.execute()
        self.assertEqual(expected_keywords, document.keywords)

    def test_create_document_from_template_updates_docproperties(self):
        expected_title = 'My title'
        self.login(self.regular_user)
        template = create(Builder('document')
                          .within(self.dossiertemplate)
                          .with_asset_file('with_gever_properties.docx'))

        with TemporaryDocFile(template.file) as tmpfile:
            template_properties = CustomProperties(Document(tmpfile.path)).items()
        self.assertItemsEqual(
            [('ogg.dossier.reference_number', 'Client1 / 2')],
            template_properties)

        command = CreateDocumentFromTemplateCommand(
            self.dossier, template, expected_title)
        document = command.execute()

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path))

        self.assertEquals(
            'Client1 1.1 / 1',
            properties['ogg.dossier.reference_number'])
        self.assertEquals(
            u'B\xe4rfuss K\xe4thi',
            properties['ogg.user.title'])

    def test_docproperties_is_updated_before_the_initial_version(self):
        self.login(self.regular_user)
        template = create(Builder('document')
                          .within(self.dossiertemplate)
                          .with_asset_file('with_gever_properties.docx'))

        with TemporaryDocFile(template.file) as tmpfile:
            template_properties = CustomProperties(Document(tmpfile.path)).items()

        self.assertItemsEqual(
            [('ogg.dossier.reference_number', 'Client1 / 2')],
            template_properties)

        command = CreateDocumentFromTemplateCommand(self.dossier, template, 'My title')
        document = command.execute()
        versioner = Versioner(document)

        self.assertFalse(versioner.has_initial_version())
        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path))

        self.assertEquals('Client1 1.1 / 1',
                          properties['ogg.dossier.reference_number'])


class TestCreateDossierFromTemplateCommand(IntegrationTestCase):

    def test_create_dossier_from_template(self):
        self.login(self.regular_user)
        command = CreateDossierFromTemplateCommand(self.dossier,
                                                   self.dossiertemplate,
                                                   responsible=self.regular_user.id)
        dossier = command.execute()
        self.assertEqual(self.dossiertemplate.title, dossier.title)
        self.assertTrue(IDossierMarker.providedBy(dossier))
        self.assertEqual(self.regular_user.id, IDossier(dossier).responsible)
