from datetime import date
from datetime import datetime
from datetime import timedelta
from docx import Document
from docxcompose.properties import CustomProperties
from ftw.builder import Builder
from ftw.builder import create
from ftw.testbrowser import browsing
from ftw.testing import freeze
from opengever.base.role_assignments import RoleAssignmentManager
from opengever.base.role_assignments import SharingRoleAssignment
from opengever.base.utils import get_date_with_delta_excluding_weekends
from opengever.core.testing import toggle_feature
from opengever.document.docprops import TemporaryDocFile
from opengever.dossier.behaviors.dossier import IDossier
from opengever.dossier.dossiertemplate.behaviors import IDossierTemplate
from opengever.dossier.dossiertemplate.interfaces import IDossierTemplateSettings
from opengever.kub.testing import KuBIntegrationTestCase
from opengever.ogds.base.actor import INTERACTIVE_ACTOR_CURRENT_USER_ID
from opengever.ogds.base.actor import INTERACTIVE_ACTOR_RESPONSIBLE_ID
from opengever.ogds.models.team import Team
from opengever.tasktemplates.interfaces import IContainParallelProcess
from opengever.tasktemplates.interfaces import IContainSequentialProcess
from opengever.tasktemplates.interfaces import IPartOfParallelProcess
from opengever.tasktemplates.interfaces import IPartOfSequentialProcess
from opengever.testing import IntegrationTestCase
from plone import api
from zExceptions import BadRequest
from zope.component import getUtility
from zope.schema.interfaces import IVocabularyFactory
import json
import pytz
import requests_mock


def flatten_tree(obj):
    flattened_items = []

    def flatten(items):
        for item in items:
            flattened_items.append(item)
            flatten(item.get('items', []))

    flatten([obj])
    return flattened_items


class TestDocumentFromTemplatePost(IntegrationTestCase):

    @browsing
    def test_creates_document_from_template_within_dossier(self, browser):
        self.login(self.regular_user, browser)

        browser.open(
            '{}/@vocabularies/opengever.dossier.DocumentTemplatesVocabulary'.format(
                self.portal.absolute_url()),
            headers=self.api_headers)
        template = browser.json['items'][0]

        data = {'template': template,
                'title': u'New d\xf6cument'}

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()

        self.assertEqual(u'New d\xf6cument', document.title)
        self.assertEqual(date.today(), document.document_date)

    @browsing
    def test_creates_document_from_template_within_task(self, browser):
        self.login(self.regular_user, browser)

        browser.open(
            '{}/@vocabularies/opengever.dossier.DocumentTemplatesVocabulary'.format(
                self.portal.absolute_url()),
            headers=self.api_headers)
        template = browser.json['items'][0]

        data = {'template': template,
                'title': u'New d\xf6cument'}

        with self.observe_children(self.task) as children:
            browser.open('{}/@document-from-template'.format(
                         self.task.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()

        self.assertEqual(u'New d\xf6cument', document.title)
        self.assertEqual(date.today(), document.document_date)

    @browsing
    def test_creates_document_from_template_within_forwarding(self, browser):
        self.login(self.secretariat_user, browser)

        browser.open(
            '{}/@vocabularies/opengever.dossier.DocumentTemplatesVocabulary'.format(
                self.portal.absolute_url()),
            headers=self.api_headers)
        template = browser.json['items'][0]

        data = {'template': template,
                'title': u'New d\xf6cument'}

        with self.observe_children(self.inbox_forwarding) as children:
            browser.open('{}/@document-from-template'.format(
                         self.inbox_forwarding.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()

        self.assertEqual(u'New d\xf6cument', document.title)
        self.assertEqual(date.today(), document.document_date)

    @browsing
    def test_requires_title(self, browser):
        self.login(self.regular_user, browser)

        browser.open(
            '{}/@vocabularies/opengever.dossier.DocumentTemplatesVocabulary'.format(
                self.portal.absolute_url()),
            headers=self.api_headers)
        template = browser.json['items'][0]

        data = {'template': template}

        browser.exception_bubbling = True
        with self.assertRaises(BadRequest) as exc:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual('Missing parameter title', str(exc.exception))

    @browsing
    def test_requires_template(self, browser):
        self.login(self.regular_user, browser)

        data = {'title': u'New d\xf6cument'}

        browser.exception_bubbling = True
        with self.assertRaises(BadRequest) as exc:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual('Missing parameter template', str(exc.exception))

    @browsing
    def test_requires_add_permission(self, browser):
        with self.login(self.regular_user):
            RoleAssignmentManager(self.portal).add_or_update_assignment(
                SharingRoleAssignment(self.reader_user.getId(), ['Reader']),
            )
            RoleAssignmentManager(self.dossier).add_or_update_assignment(
                SharingRoleAssignment(self.reader_user.getId(), ['Reader']),
            )

        self.login(self.reader_user, browser)
        data = {'template': {'token': '1234567890'},
                'title': u'New d\xf6cument'}

        with browser.expect_unauthorized():
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

    @browsing
    def test_raises_bad_request_if_recipient_is_passed_and_kub_disabled(self, browser):
        self.login(self.regular_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'recipient': KuBIntegrationTestCase.person_jean}

        browser.exception_bubbling = True
        with self.assertRaises(BadRequest) as exc:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            'recipient, sender and participations are only supported when KuB feature is active',
            str(exc.exception))

    @browsing
    def test_raises_bad_request_if_sender_is_passed_and_kub_disabled(self, browser):
        self.login(self.regular_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'sender': KuBIntegrationTestCase.person_jean}

        browser.exception_bubbling = True
        with self.assertRaises(BadRequest) as exc:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            'recipient, sender and participations are only supported when KuB feature is active',
            str(exc.exception))

    @browsing
    def test_raises_bad_request_if_participations_is_passed_and_kub_disabled(self, browser):
        self.login(self.regular_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'participations': [
                    {'participant_id': KuBIntegrationTestCase.person_jean,
                     'role': 'final-drawing'}]}

        browser.exception_bubbling = True
        with self.assertRaises(BadRequest) as exc:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            'recipient, sender and participations are only supported when KuB feature is active',
            str(exc.exception))


@requests_mock.Mocker()
class TestDocumentFromTemplatePostWithKubFeatureEnabled(KuBIntegrationTestCase):

    features = (
        'doc-properties',
    )

    document_date = datetime(2015, 9, 28, 0, 0)

    expected_doc_properties = [
        ('Document.ReferenceNumber', 'Client1 1.1 / 1 / 46'),
        ('Document.SequenceNumber', '46'),
        ('Dossier.ReferenceNumber', 'Client1 1.1 / 1'),
        ('Dossier.Title', u'Vertr\xe4ge mit der kantonalen Finanzverwaltung'),
        ('Test', 'Peter'),
        ('User.FullName', u'K\xf6nig J\xfcrgen'),
        ('User.ID', 'jurgen.konig'),
        ('ogg.document.creator.user.email', 'jurgen.konig@gever.local'),
        ('ogg.document.creator.user.firstname', u'J\xfcrgen'),
        ('ogg.document.creator.user.lastname', u'K\xf6nig'),
        ('ogg.document.creator.user.title', u'K\xf6nig J\xfcrgen'),
        ('ogg.document.creator.user.userid', 'jurgen.konig'),
        ('ogg.document.document_date', document_date),
        ('ogg.document.reference_number', 'Client1 1.1 / 1 / 46'),
        ('ogg.document.sequence_number', '46'),
        ('ogg.document.title', u'New d\xf6cument'),
        ('ogg.document.version_number', 0),
        ('ogg.dossier.external_reference', u'qpr-900-9001-\xf7'),
        ('ogg.dossier.reference_number', 'Client1 1.1 / 1'),
        ('ogg.dossier.sequence_number', '1'),
        ('ogg.dossier.title', u'Vertr\xe4ge mit der kantonalen Finanzverwaltung'),
        ('ogg.user.email', 'jurgen.konig@gever.local'),
        ('ogg.user.firstname', u'J\xfcrgen'),
        ('ogg.user.lastname', u'K\xf6nig'),
        ('ogg.user.title', u'K\xf6nig J\xfcrgen'),
        ('ogg.user.userid', 'jurgen.konig')]

    @browsing
    def test_creates_document_from_template_without_recipient(self, mocker, browser):
        self.login(self.secretariat_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument'}

        with freeze(self.document_date), self.observe_children(self.dossier) as children:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()
        self.assertEqual(u'New d\xf6cument', document.title)

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()

        self.assertItemsEqual(self.expected_doc_properties, properties)

    @browsing
    def test_creates_document_from_template_with_kub_recipient(self, mocker, browser):
        self.login(self.secretariat_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'recipient': self.person_jean}

        self.mock_get_by_id(mocker, self.person_jean)
        with freeze(self.document_date), self.observe_children(self.dossier) as children:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()
        self.assertEqual(u'New d\xf6cument', document.title)

        expected_doc_properties = self.expected_doc_properties + [
            ('ogg.recipient.contact.description', u''),
            ('ogg.recipient.email.address', 'Jean.dupon@example.com'),
            ('ogg.recipient.person.academic_title', u''),
            ('ogg.recipient.address.block', 'Herr\nJean Dupont\nTeststrasse 43\n9999 Bern'),
            ('ogg.recipient.address.extra_line_1', u''),
            ('ogg.recipient.address.extra_line_2', u''),
            ('ogg.recipient.address.zip_code', '9999'),
            ('ogg.recipient.address.city', 'Bern'),
            ('ogg.recipient.contact.title', 'Dupont Jean'),
            ('ogg.recipient.address.street', 'Teststrasse 43'),
            ('ogg.recipient.address.country', 'Schweiz'),
            ('ogg.recipient.person.date_of_birth', datetime(1992, 5, 15, 0, 0)),
            ('ogg.recipient.person.firstname', 'Jean'),
            ('ogg.recipient.person.lastname', 'Dupont'),
            ('ogg.recipient.person.salutation', 'Herr'),
            ('ogg.recipient.phone.number', '666 666 66 66')]

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()

        self.assertItemsEqual(expected_doc_properties, properties)

    @browsing
    def test_creates_document_from_template_with_ogds_recipient(self, mocker, browser):
        self.login(self.secretariat_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'recipient': self.regular_user.getId()}

        with freeze(self.document_date), self.observe_children(self.dossier) as children:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()
        self.assertEqual(u'New d\xf6cument', document.title)

        expected_doc_properties = self.expected_doc_properties + [
            ('ogg.recipient.address.block', u'Frau\nK\xe4thi B\xe4rfuss\nKappelenweg 13\n1234 Vorkappelen'),
            ('ogg.recipient.address.city', 'Vorkappelen'),
            ('ogg.recipient.address.country', 'Schweiz'),
            ('ogg.recipient.address.street', 'Kappelenweg 13, Postfach 1234'),
            ('ogg.recipient.address.zip_code', '1234'),
            ('ogg.recipient.contact.description', 'nix'),
            ('ogg.recipient.contact.title', u'B\xe4rfuss K\xe4thi'),
            ('ogg.recipient.email.address', 'foo@example.com'),
            ('ogg.recipient.person.firstname', u'K\xe4thi'),
            ('ogg.recipient.person.lastname', u'B\xe4rfuss'),
            ('ogg.recipient.person.salutation', 'Frau'),
            ('ogg.recipient.phone.number', '012 34 56 78'),
            ('ogg.recipient.url.url', 'http://www.example.com'),
        ]

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
        self.assertItemsEqual(expected_doc_properties, properties)

    @browsing
    def test_creates_document_from_template_with_kub_sender(self, mocker, browser):
        self.login(self.secretariat_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'sender': self.person_jean}

        self.mock_get_by_id(mocker, self.person_jean)
        with freeze(self.document_date), self.observe_children(self.dossier) as children:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()
        self.assertEqual(u'New d\xf6cument', document.title)

        expected_doc_properties = self.expected_doc_properties + [
            ('ogg.sender.contact.description', u''),
            ('ogg.sender.email.address', 'Jean.dupon@example.com'),
            ('ogg.sender.person.academic_title', u''),
            ('ogg.sender.address.block', 'Herr\nJean Dupont\nTeststrasse 43\n9999 Bern'),
            ('ogg.sender.address.extra_line_1', u''),
            ('ogg.sender.address.extra_line_2', u''),
            ('ogg.sender.address.zip_code', '9999'),
            ('ogg.sender.address.city', 'Bern'),
            ('ogg.sender.contact.title', 'Dupont Jean'),
            ('ogg.sender.address.street', 'Teststrasse 43'),
            ('ogg.sender.address.country', 'Schweiz'),
            ('ogg.sender.person.date_of_birth', datetime(1992, 5, 15, 0, 0)),
            ('ogg.sender.person.firstname', 'Jean'),
            ('ogg.sender.person.lastname', 'Dupont'),
            ('ogg.sender.person.salutation', 'Herr'),
            ('ogg.sender.phone.number', '666 666 66 66')]

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()

        self.assertItemsEqual(expected_doc_properties, properties)

    @browsing
    def test_creates_document_from_template_with_kub_membership_sender(self, mocker, browser):
        self.login(self.secretariat_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'sender': self.memb_jean_ftw}

        self.mock_get_by_id(mocker, self.memb_jean_ftw)
        with freeze(self.document_date), self.observe_children(self.dossier) as children:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()
        self.assertEqual(u'New d\xf6cument', document.title)

        expected_doc_properties = self.expected_doc_properties + [
            ('ogg.sender.address.block', '4Teamwork\nHerr Jean Dupont\nc/o John Doe\nDammweg 9\n3013 Bern'),
            ('ogg.sender.address.city', 'Bern'),
            ('ogg.sender.address.country', 'Schweiz'),
            ('ogg.sender.address.extra_line_1', 'c/o John Doe'),
            ('ogg.sender.address.extra_line_2', u''),
            ('ogg.sender.address.street', 'Dammweg 9'),
            ('ogg.sender.address.zip_code', '3013'),
            ('ogg.sender.contact.description', u''),
            ('ogg.sender.contact.title', 'Dupont Jean - 4Teamwork (CEO)'),
            ('ogg.sender.email.address', 'Jean.dupon@example.com'),
            ('ogg.sender.organization.name', '4Teamwork'),
            ('ogg.sender.organization.phone.number', '111 111 11 11'),
            ('ogg.sender.orgrole.department', u''),
            ('ogg.sender.orgrole.description', u''),
            ('ogg.sender.orgrole.function', 'CEO'),
            ('ogg.sender.person.academic_title', u''),
            ('ogg.sender.person.date_of_birth', datetime(1992, 5, 15, 0, 0)),
            ('ogg.sender.person.firstname', 'Jean'),
            ('ogg.sender.person.lastname', 'Dupont'),
            ('ogg.sender.person.salutation', 'Herr'),
            ('ogg.sender.phone.number', '666 666 66 66')]

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()

        self.assertItemsEqual(expected_doc_properties, properties)

    @browsing
    def test_creates_document_from_template_with_ogds_sender(self, mocker, browser):
        self.login(self.secretariat_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'sender': self.regular_user.getId()}

        with freeze(self.document_date), self.observe_children(self.dossier) as children:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()
        self.assertEqual(u'New d\xf6cument', document.title)

        expected_doc_properties = self.expected_doc_properties + [
            ('ogg.sender.address.block', u'Frau\nK\xe4thi B\xe4rfuss\nKappelenweg 13\n1234 Vorkappelen'),
            ('ogg.sender.address.city', 'Vorkappelen'),
            ('ogg.sender.address.country', 'Schweiz'),
            ('ogg.sender.address.street', 'Kappelenweg 13, Postfach 1234'),
            ('ogg.sender.address.zip_code', '1234'),
            ('ogg.sender.contact.description', 'nix'),
            ('ogg.sender.contact.title', u'B\xe4rfuss K\xe4thi'),
            ('ogg.sender.email.address', 'foo@example.com'),
            ('ogg.sender.person.firstname', u'K\xe4thi'),
            ('ogg.sender.person.lastname', u'B\xe4rfuss'),
            ('ogg.sender.person.salutation', 'Frau'),
            ('ogg.sender.phone.number', '012 34 56 78'),
            ('ogg.sender.url.url', 'http://www.example.com'),
        ]

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
        self.assertItemsEqual(expected_doc_properties, properties)

    @browsing
    def test_creates_document_from_template_with_participations(self, mocker, browser):
        self.login(self.secretariat_user, browser)

        data = {'template': self.docprops_template.UID(),
                'title': u'New d\xf6cument',
                'participations': [
                    {'participant_id': self.person_jean, 'role': 'final-drawing'},
                    {'participant_id': self.regular_user.getId(), 'role': 'participation'}]}

        self.mock_get_by_id(mocker, self.person_jean)
        with freeze(self.document_date), self.observe_children(self.dossier) as children:
            browser.open('{}/@document-from-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        document = children['added'].pop()
        self.assertEqual(u'New d\xf6cument', document.title)

        expected_doc_properties = self.expected_doc_properties + [
            ('ogg.final-drawing.address.block',
             'Herr\nJean Dupont\nTeststrasse 43\n9999 Bern'),
            ('ogg.final-drawing.address.city', 'Bern'),
            ('ogg.final-drawing.address.country', 'Schweiz'),
            ('ogg.final-drawing.address.extra_line_1', u''),
            ('ogg.final-drawing.address.extra_line_2', u''),
            ('ogg.final-drawing.address.street', 'Teststrasse 43'),
            ('ogg.final-drawing.address.zip_code', '9999'),
            ('ogg.final-drawing.contact.description', u''),
            ('ogg.final-drawing.contact.title', 'Dupont Jean'),
            ('ogg.final-drawing.email.address', 'Jean.dupon@example.com'),
            ('ogg.final-drawing.person.academic_title', u''),
            ('ogg.final-drawing.person.date_of_birth', datetime(1992, 5, 15, 0, 0)),
            ('ogg.final-drawing.person.firstname', 'Jean'),
            ('ogg.final-drawing.person.lastname', 'Dupont'),
            ('ogg.final-drawing.person.salutation', 'Herr'),
            ('ogg.final-drawing.phone.number', '666 666 66 66'),
            ('ogg.participation.address.block',
             u'Frau\nK\xe4thi B\xe4rfuss\nKappelenweg 13\n1234 Vorkappelen'),
            ('ogg.participation.address.city', 'Vorkappelen'),
            ('ogg.participation.address.country', 'Schweiz'),
            ('ogg.participation.address.street', 'Kappelenweg 13, Postfach 1234'),
            ('ogg.participation.address.zip_code', '1234'),
            ('ogg.participation.contact.description', 'nix'),
            ('ogg.participation.contact.title', u'B\xe4rfuss K\xe4thi'),
            ('ogg.participation.email.address', 'foo@example.com'),
            ('ogg.participation.person.firstname', u'K\xe4thi'),
            ('ogg.participation.person.lastname', u'B\xe4rfuss'),
            ('ogg.participation.person.salutation', 'Frau'),
            ('ogg.participation.phone.number', '012 34 56 78'),
            ('ogg.participation.url.url', 'http://www.example.com')]

        with TemporaryDocFile(document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
        self.assertItemsEqual(expected_doc_properties, properties)


class TestDossierFromTemplatePost(IntegrationTestCase):

    features = ('dossiertemplate', )

    @browsing
    def test_create_dossier_from_template_within_repository_folder(self, browser):
        self.login(self.regular_user, browser)
        browser.open(self.leaf_repofolder,
                     view='@vocabularies/opengever.dossier.DossierTemplatesVocabulary',
                     headers=self.api_headers)
        template = browser.json['items'][0]

        data = {'template': template,
                'title': u'New d\xf6ssier',
                'responsible': self.regular_user.getId()}

        with self.observe_children(self.leaf_repofolder) as children:
            browser.open('{}/@dossier-from-template'.format(
                         self.leaf_repofolder.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        dossier = children['added'].pop()
        self.assertEqual(u'New d\xf6ssier', dossier.title)
        self.assertEqual(self.regular_user.getId(), IDossier(dossier).responsible)
        self.assertEqual([u'Werkst\xe4tte', u'Anfragen'],
                         [obj.title for obj in dossier.listFolderContents()])

        subdossier = dossier.listFolderContents()[1]
        self.assertEqual([u'Baumsch\xfctze'],
                         [obj.title for obj in subdossier.listFolderContents()])
        self.assertEqual(self.regular_user.getId(), IDossier(subdossier).responsible)

    @browsing
    def test_create_dossier_from_template_properly_sets_role_manager_local_roles(self, browser):
        self.activate_feature('grant_role_manager_to_responsible')
        self.login(self.regular_user, browser)
        browser.open(self.leaf_repofolder,
                     view='@vocabularies/opengever.dossier.DossierTemplatesVocabulary',
                     headers=self.api_headers)
        template = browser.json['items'][0]

        data = {'template': template,
                'title': u'New d\xf6ssier',
                'responsible': self.regular_user.getId()}

        with self.observe_children(self.leaf_repofolder) as children:
            browser.open('{}/@dossier-from-template'.format(
                         self.leaf_repofolder.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        dossier = children['added'].pop()
        self.assertEqual(self.regular_user.getId(), IDossier(dossier).responsible)
        self.assertEqual(((self.regular_user.getId(), ('Role Manager', 'Owner')), ),
                         dossier.get_local_roles())

        subdossier = dossier.listFolderContents()[1]
        self.assertEqual(self.regular_user.id, IDossier(subdossier).responsible)
        self.assertEqual(((self.regular_user.id, ('Role Manager', 'Owner')), ),
                         subdossier.get_local_roles())

    @browsing
    def test_raise_unauthorized_if_dossiertemplate_feature_is_not_available(self, browser):
        self.login(self.regular_user, browser)
        toggle_feature(IDossierTemplateSettings, enabled=False)

        browser.open(self.leaf_repofolder,
                     view='@vocabularies/opengever.dossier.DossierTemplatesVocabulary',
                     headers=self.api_headers)
        template = browser.json['items'][0]

        data = {'template': template,
                'title': u'New d\xf6ssier',
                'responsible': self.regular_user.getId()}

        with browser.expect_unauthorized():
            browser.open('{}/@dossier-from-template'.format(
                         self.leaf_repofolder.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

    @browsing
    def test_invalid_keyword_raises_bad_request(self, browser):
        self.login(self.regular_user, browser)
        self.dossiertemplate.restrict_keywords = True
        vocabulary = getUtility(IVocabularyFactory,
                                name='opengever.dossier.DossierTemplatesVocabulary')
        token = vocabulary(self.leaf_repofolder).getTerm(self.dossiertemplate).token
        data = {'template': {'token': token},
                'title': u'New d\xf6ssier',
                'responsible': self.regular_user.getId(),
                'keywords': ['secret', 'invalid']}

        with browser.expect_http_error(400):
            browser.open('{}/@dossier-from-template'.format(
                         self.leaf_repofolder.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {"message": "Keyword 'invalid' is not allowed",
             "type": "BadRequest"},
            browser.json)

    @browsing
    def test_missing_template_raises_bad_request(self, browser):
        self.login(self.regular_user, browser)
        data = {'title': u'New d\xf6ssier',
                'responsible': self.regular_user.getId()}

        with browser.expect_http_error(400):
            browser.open('{}/@dossier-from-template'.format(
                         self.leaf_repofolder.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {"message": "Missing parameter template",
             "type": "BadRequest"},
            browser.json)

    @browsing
    def test_invalid_template_raises_bad_request(self, browser):
        self.login(self.administrator, browser)
        template = create(Builder("dossiertemplate")
                          .within(self.templates)
                          .titled(u"Bauvorhaben gross"))
        vocabulary = getUtility(IVocabularyFactory,
                                name='opengever.dossier.DossierTemplatesVocabulary')
        token = vocabulary(self.leaf_repofolder).getTerm(self.dossiertemplate).token
        self.set_related_items(
            self.leaf_repofolder, [template, ],
            fieldname='addable_dossier_templates')

        self.login(self.regular_user, browser)

        data = {'template': {'token': token},
                'title': u'New d\xf6ssier',
                'responsible': self.regular_user.getId()}
        with browser.expect_http_error(400):
            browser.open('{}/@dossier-from-template'.format(
                         self.leaf_repofolder.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {"message": "Invalid token '{}'".format(token),
             "type": "BadRequest"},
            browser.json)

    @browsing
    def test_create_dossier_from_template_with_related_documents(self, browser):
        self.login(self.administrator, browser)

        self.set_related_items(
            IDossierTemplate(self.dossiertemplate), [self.empty_template, ],
            fieldname='related_documents')

        self.set_related_items(
            IDossierTemplate(self.subdossiertemplate), [self.normal_template, ],
            fieldname='related_documents')

        data = {'template': { 'token': self.dossiertemplate.UID()},
                'title': u'New d\xf6ssier',
                'responsible': self.regular_user.getId()}

        with self.observe_children(self.leaf_repofolder) as children:
            browser.open('{}/@dossier-from-template'.format(
                         self.leaf_repofolder.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        dossier = children['added'].pop()

        self.assertEqual([
            u'T\xc3\xb6mpl\xc3\xb6te Leer', u'Werkst\xe4tte', u'Anfragen'],
            [obj.title for obj in dossier.listFolderContents()])

        subdossier = dossier.listFolderContents()[-1]
        self.assertEqual([u'T\xc3\xb6mpl\xc3\xb6te Normal', u'Baumsch\xfctze'],
                         [obj.title for obj in subdossier.listFolderContents()])

    @browsing
    def test_set_dossier_local_roles_from_template(self, browser):
        self.login(self.regular_user, browser)

        self.dossiertemplate.__ac_local_roles_block__ = True
        RoleAssignmentManager(self.dossiertemplate).add_or_update_assignment(
            SharingRoleAssignment(self.regular_user.getId(),
                                  ['Reader', 'Contributor', 'Editor']))
        RoleAssignmentManager(self.dossiertemplate).add_or_update_assignment(
            SharingRoleAssignment(self.dossier_manager.getId(),
                                  ['Reader']))

        self.subdossiertemplate.__ac_local_roles_block__ = True
        RoleAssignmentManager(self.subdossiertemplate).add_or_update_assignment(
            SharingRoleAssignment(self.regular_user.getId(),
                                  ['Reader', 'Contributor', 'Editor']))
        RoleAssignmentManager(self.subdossiertemplate).add_or_update_assignment(
            SharingRoleAssignment(self.dossier_manager.getId(),
                                  ['Reader', 'Editor']))

        browser.open(self.leaf_repofolder,
                     view='@vocabularies/opengever.dossier.DossierTemplatesVocabulary',
                     headers=self.api_headers)
        template = browser.json['items'][0]

        data = {'template': template,
                'title': u'New d\xf6ssier',
                'responsible': self.regular_user.getId()}

        with self.observe_children(self.leaf_repofolder) as children:
            browser.open('{}/@dossier-from-template'.format(
                         self.leaf_repofolder.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        dossier = children['added'].pop()
        self.assertTrue(dossier.__ac_local_roles_block__)
        self.assertEqual(
            [
                {
                    'cause': 3,
                    'roles': ['Reader', 'Contributor', 'Editor'],
                    'reference': None,
                    'principal': 'regular_user',
                },
                {
                    'cause': 3,
                    'roles': ['Reader'],
                    'reference': None,
                    'principal': 'dossier_manager',
                }
            ], RoleAssignmentManager(dossier).storage._storage()
        )

        subdossier = dossier.listFolderContents()[1]
        self.assertTrue(subdossier.__ac_local_roles_block__)
        self.assertEqual(
            [
                {
                    'cause': 3,
                    'roles': ['Reader', 'Contributor', 'Editor'],
                    'reference': None,
                    'principal': 'regular_user',
                },
                {
                    'cause': 3,
                    'roles': ['Reader', 'Editor'],
                    'reference': None,
                    'principal': 'dossier_manager',
                }
            ], RoleAssignmentManager(subdossier).storage._storage()
        )


class TestTriggerTaskTemplatePost(IntegrationTestCase):

    def _get_task_template_item(self, browser, title=None):
        """Return a task template folder vocabulary item.

        Return the first task template folder item from the vocabulary if no
        title is given, otherwise return the specified item.
        """
        browser.open(
            '{}/@vocabularies/opengever.tasktemplates.active_tasktemplatefolders'.format(
                self.portal.absolute_url()),
            headers=self.api_headers)

        if title:
            for item in browser.json['items']:
                if item['title'] == title:
                    return item
            raise ValueError(u'no item with title {} found'.format(title))

        return browser.json['items'][0]

    @browsing
    def test_trigger_with_minimal_required_input(self, browser):
        self.login(self.regular_user, browser)

        folder_data = self._get_task_template_item(browser)
        folder_data.pop('title')

        data = {
            'tasktemplatefolder': folder_data,
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'start_immediately': True
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        main_task = children['added'].pop()

        self.assertEqual(browser.json['@id'], main_task.absolute_url())

        self.assertEqual(u'Verfahren Neuanstellung', main_task.title)
        self.assertEqual(self.regular_user.getId(), main_task.issuer)
        self.assertEqual(self.regular_user.getId(), main_task.responsible)
        self.assertEqual('fa', main_task.responsible_client)
        self.assertEqual('direct-execution', main_task.task_type)
        self.assertEqual('task-state-in-progress',
                         api.content.get_state(main_task))

        subtasks = main_task.listFolderContents()
        self.assertEqual(1, len(subtasks))
        subtask = subtasks.pop()
        self.assertEqual(u'Arbeitsplatz einrichten.', subtask.title)
        self.assertEqual('robert.ziegler', subtask.issuer)
        self.assertEqual('robert.ziegler', subtask.responsible)
        self.assertEqual('fa', subtask.responsible_client)
        self.assertEqual('task-state-open',
                         api.content.get_state(subtask))

    @browsing
    def test_overriding_main_task_title_text_and_deadline(self, browser):
        self.login(self.regular_user, browser)

        folder_data = self._get_task_template_item(browser)
        folder_data.pop('title')

        data = {
            'tasktemplatefolder': folder_data,
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'deadline': u'2021-12-12',
            'title': u'Neuanstellung Hugo B\xf6ss',
            'text': u'Bla bla',
            'start_immediately': True
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        main_task = children['added'].pop()

        self.assertEqual(u'Neuanstellung Hugo B\xf6ss', main_task.title)
        self.assertEqual(u'Bla bla', main_task.text)
        self.assertEqual(date(2021, 12, 12), main_task.deadline)

    @browsing
    def test_overriding_sub_task_title_text_and_deadline(self, browser):
        self.login(self.regular_user, browser)

        folder_data = self._get_task_template_item(browser)
        folder_data.pop('title')

        data = {
            'tasktemplatefolder': folder_data,
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                    'deadline': u'2021-12-12',
                    'title': u'Arbeitsplatz Hugo B\xf6ss',
                    'text': None,
                }
            ],
            'start_immediately': True
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        main_task = children['added'].pop()
        subtasks = main_task.listFolderContents()
        self.assertEqual(1, len(subtasks))
        subtask = subtasks.pop()

        self.assertEqual(u'Arbeitsplatz Hugo B\xf6ss', subtask.title)
        self.assertIsNone(subtask.text)
        self.assertEqual(date(2021, 12, 12), subtask.deadline)

    @browsing
    def test_invalid_main_task_title_raises_bad_request(self, browser):
        self.login(self.regular_user, browser)

        folder_data = self._get_task_template_item(browser)
        folder_data.pop('title')

        data = {
            'tasktemplatefolder': folder_data,
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'title': None,
            'text': u'Bla bla',
            'start_immediately': True
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {u'message': u'Invalid title "None"', u"type": u"BadRequest"},
            browser.json)

    @browsing
    def test_invalid_sub_task_title_raises_bad_request(self, browser):
        self.login(self.regular_user, browser)

        folder_data = self._get_task_template_item(browser)
        folder_data.pop('title')

        data = {
            'tasktemplatefolder': folder_data,
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                    'title': '',
                    'text': u'Bla bla',
                }
            ],
            'title': 'Valid',
            'text': u'Bla bla',
            'start_immediately': True
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {u'message': u'Invalid title ""', u"type": u"BadRequest"},
            browser.json)

    @browsing
    def test_invalid_task_deadline_raises_bad_request(self, browser):
        self.login(self.regular_user, browser)

        folder_data = self._get_task_template_item(browser)

        data = {
            'tasktemplatefolder': folder_data,
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'title': 'A title',
            'text': u'Bla bla',
            'deadline': None,
            'start_immediately': True
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {u'message': u'Invalid deadline "None"', u"type": u"BadRequest"},
            browser.json)

    @browsing
    def test_trigger_with_non_nested_task_template_folder_input(self, browser):
        self.login(self.regular_user, browser)

        folder_data = self._get_task_template_item(browser)
        folder_data.pop('title')

        data = {
            'tasktemplatefolder': folder_data['token'],
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'start_immediately': True
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))

    @browsing
    def test_selecting_no_tasktemplates_raises_badrequest(self, browser):
        self.login(self.regular_user, browser=browser)

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [],
            'start_immediately': True
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {"message": "At least one tasktemplate is required",
             "type": "BadRequest"},
            browser.json)

    @browsing
    def test_providing_no_tasktemplates_raises_badrequest(self, browser):
        self.login(self.regular_user, browser=browser)

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'start_immediately': True
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {"message": "At least one tasktemplate is required",
             "type": "BadRequest"},
            browser.json)

    @browsing
    def test_providing_no_templatefolder_raises_badrequest(self, browser):
        self.login(self.regular_user, browser=browser)

        data = {
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'start_immediately': True
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {"message": "[('tasktemplatefolder', RequiredMissing('tasktemplatefolder'))]",
             "type": "BadRequest"},
            browser.json)

    @browsing
    def test_omitting_start_immediately_raises_badrequest(self, browser):
        self.login(self.regular_user, browser=browser)

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ]
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {"message": "[('start_immediately', RequiredMissing('start_immediately'))]",
             "type": "BadRequest"},
            browser.json)

    @browsing
    def test_task_templates_outside_selected_templatefolder_are_invalid(self, browser):
        self.login(self.regular_user, browser=browser)

        foreign_task_template_folder = create(
            Builder('tasktemplatefolder')
            .titled(u'Geheim')
            .in_state('tasktemplatefolder-state-activ')
            .having(sequence_type='parallel')
            .within(self.templates)
        )
        foreign_task_template = create(
            Builder('tasktemplate')
            .titled(u'Sag ich nicht')
            .having(**{
                'issuer': INTERACTIVE_ACTOR_CURRENT_USER_ID,
                'responsible_client': 'fa',
                'responsible': 'robert.ziegler',
                'deadline': 10,
            })
            .within(foreign_task_template_folder)
        )

        data = {
            'tasktemplatefolder': self._get_task_template_item(
                browser, title=u'Verfahren Neuanstellung'),
            'tasktemplates': [
                {
                    '@id': foreign_task_template.absolute_url(),
                }
            ],
            'start_immediately': True
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {'message': 'The tasktemplate {} is outside the selected '
                        'template folder'.format(
                            foreign_task_template.absolute_url()),
             'type': 'BadRequest'},
            browser.json)

    @browsing
    def test_main_task_deadline_is_the_highest_template_deadline_plus_five(self, browser):
        self.login(self.regular_user, browser=browser)

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'start_immediately': True
        }
        with freeze(datetime(2018, 10, 18, 7, 33, tzinfo=pytz.utc)):
            with self.observe_children(self.dossier) as children:
                browser.open('{}/@trigger-task-template'.format(
                             self.dossier.absolute_url()),
                             data=json.dumps(data),
                             headers=self.api_headers)

            main_task = children['added'].pop()
            self.assertEqual(
                get_date_with_delta_excluding_weekends(date.today() + timedelta(days=10), 5),
                main_task.deadline)

    @browsing
    def test_create_parallel_tasks(self, browser):
        self.login(self.regular_user, browser)
        self.tasktemplatefolder.sequence_type = u'parallel'

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'start_immediately': False
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        main_task = children['added'].pop()
        self.assertTrue(IContainParallelProcess.providedBy(main_task))
        self.assertEqual('task-state-in-progress',
                         api.content.get_state(main_task))
        for subtask in main_task.listFolderContents():
            self.assertTrue(IPartOfParallelProcess.providedBy(subtask))
            self.assertEqual('task-state-open',
                             api.content.get_state(subtask))

    @browsing
    def test_create_sequential_tasks(self, browser):
        self.login(self.regular_user, browser)
        self.tasktemplatefolder.sequence_type = u'sequential'

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                }
            ],
            'start_immediately': False
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        main_task = children['added'].pop()
        self.assertTrue(IContainSequentialProcess.providedBy(main_task))
        self.assertEqual('task-state-in-progress',
                         api.content.get_state(main_task))
        for subtask in main_task.listFolderContents():
            self.assertTrue(IPartOfSequentialProcess.providedBy(subtask))
            self.assertEqual('task-state-planned',
                             api.content.get_state(subtask))

    @browsing
    def test_respects_start_immediately_flag(self, browser):
        self.login(self.regular_user, browser)
        self.tasktemplatefolder.sequence_type = u'sequential'

        template_2 = create(Builder('tasktemplate')
                            .titled(u'Noch was.')
                            .having(issuer=INTERACTIVE_ACTOR_CURRENT_USER_ID,
                                    responsible_client='fa',
                                    responsible='robert.ziegler',
                                    deadline=10,)
                            .within(self.tasktemplatefolder))

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url()
                },
                {
                    '@id': template_2.absolute_url()
                }
            ],
            'start_immediately': True
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        main_task = children['added'].pop()
        self.assertEquals('task-state-in-progress',
                          api.content.get_state(main_task))

        subtask_1, subtask_2 = main_task.listFolderContents()
        self.assertEquals('task-state-open',
                          api.content.get_state(subtask_1))
        self.assertEquals('task-state-planned',
                          api.content.get_state(subtask_2))

    @browsing
    def test_creates_subtasks_for_selected_task_templates(self, browser):
        self.login(self.regular_user, browser=browser)

        notebook = create(Builder('tasktemplate')
                          .titled(u'Notebook einrichten.')
                          .having(issuer=INTERACTIVE_ACTOR_CURRENT_USER_ID,
                                  responsible_client='fa',
                                  responsible='robert.ziegler',
                                  deadline=10)
                          .within(self.tasktemplatefolder))

        user_accounts = create(Builder('tasktemplate')
                               .titled(u'User Accounts erstellen.')
                               .having(issuer=INTERACTIVE_ACTOR_CURRENT_USER_ID,
                                       responsible_client='fa',
                                       responsible='robert.ziegler',
                                       deadline=10)
                               .within(self.tasktemplatefolder))

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': notebook.absolute_url(),
                },
                {
                    '@id': user_accounts.absolute_url(),
                }
            ],
            'start_immediately': False
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        main_task = children['added'].pop()
        self.assertEqual(2, len(main_task.listFolderContents()))

        self.assertItemsEqual(
            [u'Notebook einrichten.', u'User Accounts erstellen.'],
            [item.title for item in main_task.listFolderContents()])

    @browsing
    def test_set_relateditems_on_every_subtask_when_selected(self, browser):
        self.login(self.regular_user, browser=browser)

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url()
                }
            ],
            'related_documents': [
                {
                    '@id': self.document.absolute_url()
                }
            ],
            'start_immediately': False
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        subtask = children['added'].pop().listFolderContents()[0]
        self.assertEqual(
            [self.document],
            [relation.to_object for relation in subtask.relatedItems])

    @browsing
    def test_disallows_related_documents_outside_dossier(self, browser):
        self.login(self.regular_user, browser=browser)

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url()
                }
            ],
            'related_documents': [
                {
                    '@id': self.private_document.absolute_url()
                }
            ],
            'start_immediately': True
        }

        with browser.expect_http_error(400):
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)
        self.assertEqual(
            {"message":
                u"The following related_documents are invalid: "
                "{{u'@id': u'{}'}}".format(self.private_document.absolute_url()),
             "type": "BadRequest"},
            browser.json)

    @browsing
    def test_override_task_responsible_with_different_user(self, browser):
        self.login(self.regular_user, browser)

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                    'responsible': 'fa:{}'.format(self.dossier_manager.getId())
                }
            ],
            'start_immediately': True,
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        main_task = children['added'].pop()

        self.assertEqual(self.regular_user.getId(), main_task.issuer)
        self.assertEqual(self.regular_user.getId(), main_task.responsible)
        self.assertEqual('fa', main_task.responsible_client)

        subtask = main_task.listFolderContents().pop()
        self.assertEqual('robert.ziegler', subtask.issuer)
        self.assertEqual(self.dossier_manager.getId(), subtask.responsible)
        self.assertEqual('fa', subtask.responsible_client)

    @browsing
    def test_override_task_responsible_with_team(self, browser):
        self.login(self.regular_user, browser)

        team = Team.get_by(groupid='projekt_a')
        team_id = 'team:{}'.format(team.team_id)

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                    'responsible': team_id
                }
            ],
            'start_immediately': True,
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        main_task = children['added'].pop()

        self.assertEqual(self.regular_user.getId(), main_task.issuer)
        self.assertEqual(self.regular_user.getId(), main_task.responsible)
        self.assertEqual('fa', main_task.responsible_client)

        subtask = main_task.listFolderContents().pop()
        self.assertEqual('robert.ziegler', subtask.issuer)
        self.assertEqual(team_id, subtask.responsible)
        self.assertEqual('fa', subtask.responsible_client)

    @browsing
    def test_override_task_responsible_with_inbox(self, browser):
        self.login(self.regular_user, browser)

        inbox_id = 'inbox:fa'

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                    'responsible': inbox_id
                }
            ],
            'start_immediately': True,
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        main_task = children['added'].pop()

        self.assertEqual(self.regular_user.getId(), main_task.issuer)
        self.assertEqual(self.regular_user.getId(), main_task.responsible)
        self.assertEqual('fa', main_task.responsible_client)

        subtask = main_task.listFolderContents().pop()
        self.assertEqual('robert.ziegler', subtask.issuer)
        self.assertEqual(inbox_id, subtask.responsible)
        self.assertEqual('fa', subtask.responsible_client)

    @browsing
    def test_override_task_responsible_with_interactive_responsible(self, browser):
        self.login(self.regular_user, browser)
        IDossier(self.dossier).responsible = self.secretariat_user.getId()

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                    'responsible': INTERACTIVE_ACTOR_RESPONSIBLE_ID,
                }
            ],
            'start_immediately': True,
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        main_task = children['added'].pop()

        self.assertEqual(self.regular_user.getId(), main_task.issuer)
        self.assertEqual(self.regular_user.getId(), main_task.responsible)
        self.assertEqual('fa', main_task.responsible_client)

        subtask = main_task.listFolderContents().pop()
        self.assertEqual(self.secretariat_user.getId(), subtask.issuer)
        self.assertEqual(self.secretariat_user.getId(), subtask.responsible)
        self.assertEqual('fa', subtask.responsible_client)

    @browsing
    def test_override_task_responsible_with_interactive_current_user(self, browser):
        self.login(self.regular_user, browser)
        IDossier(self.dossier).responsible = self.secretariat_user.getId()

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url(),
                    'responsible': INTERACTIVE_ACTOR_CURRENT_USER_ID
                }
            ],
            'start_immediately': True,
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        self.assertEqual(1, len(children['added']))
        main_task = children['added'].pop()

        self.assertEqual(self.regular_user.getId(), main_task.issuer)
        self.assertEqual(self.regular_user.getId(), main_task.responsible)
        self.assertEqual('fa', main_task.responsible_client)

        subtask = main_task.listFolderContents().pop()
        self.assertEqual('jurgen.konig', subtask.issuer)
        self.assertEqual(self.regular_user.getId(), subtask.responsible)
        self.assertEqual('fa', subtask.responsible_client)

    @browsing
    def test_sets_task_templatefolder_predecessor(self, browser):
        self.login(self.regular_user, browser=browser)
        self.tasktemplatefolder.sequence_type = u'sequential'

        template_2 = create(Builder('tasktemplate')
                            .titled(u'Notebook einrichten.')
                            .having(issuer=INTERACTIVE_ACTOR_CURRENT_USER_ID,
                                    responsible_client='fa',
                                    responsible='robert.ziegler',
                                    deadline=10,)
                            .within(self.tasktemplatefolder))

        data = {
            'tasktemplatefolder': self._get_task_template_item(browser),
            'tasktemplates': [
                {
                    '@id': self.tasktemplate.absolute_url()
                },
                {
                    '@id': template_2.absolute_url()
                }
            ],
            'start_immediately': True
        }

        with self.observe_children(self.dossier) as children:
            browser.open('{}/@trigger-task-template'.format(
                         self.dossier.absolute_url()),
                         data=json.dumps(data),
                         headers=self.api_headers)

        main_task = children['added'].pop()
        task1, task2 = main_task.listFolderContents()

        self.assertEquals(task1.get_sql_object(),
                          task2.get_sql_object().get_previous_task())


class TestTaskTemplateStructure(IntegrationTestCase):

    @browsing
    def test_returns_recursively_the_full_serialized_objects_of_all_tasktemplates_and_folders(self, browser):
        self.activate_feature('tasktemplatefolder_nesting')

        self.login(self.administrator, browser=browser)

        subtasktemplatefolder = create(
            Builder('tasktemplatefolder')
            .titled(u'sub-tasktemplate folder 1')
            .having(sequence_type='parallel')
            .within(self.tasktemplatefolder)
        )

        create(
            Builder('tasktemplate')
            .titled(u'sub-task 1.1')
            .having(issuer=INTERACTIVE_ACTOR_CURRENT_USER_ID,
                    responsible_client='fa',
                    responsible='robert.ziegler',)
            .within(subtasktemplatefolder)
        )

        browser.open(
            '{}/@task-template-structure'.format(
                self.tasktemplatefolder.absolute_url()),
            headers=self.api_headers)

        self.assertEqual(
            [
                u'Verfahren Neuanstellung',
                u'Arbeitsplatz einrichten.',
                u'sub-tasktemplate folder 1',
                u'sub-task 1.1'
            ],
            [item.get('title') for item in flatten_tree(browser.json)]
        )

    @browsing
    def test_tasktemplates_are_ordered_by_position_in_parent(self, browser):
        self.activate_feature('tasktemplatefolder_nesting')

        self.login(self.administrator, browser=browser)

        subtasktemplatefolder = create(
            Builder('tasktemplatefolder')
            .titled(u'sub-tasktemplate folder 1')
            .having(sequence_type='parallel')
            .within(self.tasktemplatefolder)
        )

        subtasktemplate = create(
            Builder('tasktemplate')
            .titled(u'sub-task 1.1')
            .having(issuer=INTERACTIVE_ACTOR_CURRENT_USER_ID,
                    responsible_client='fa',
                    responsible='robert.ziegler',)
            .within(subtasktemplatefolder)
        )

        create(
            Builder('tasktemplate')
            .titled(u'sub-task 1.2')
            .having(issuer=INTERACTIVE_ACTOR_CURRENT_USER_ID,
                    responsible_client='fa',
                    responsible='robert.ziegler',)
            .within(subtasktemplatefolder)
        )

        browser.open(
            '{}/@task-template-structure'.format(
                self.tasktemplatefolder.absolute_url()),
            headers=self.api_headers)

        self.assertEqual(
            [
                u'Verfahren Neuanstellung',
                u'Arbeitsplatz einrichten.',
                u'sub-tasktemplate folder 1',
                u'sub-task 1.1',
                u'sub-task 1.2',
            ],
            [item.get('title') for item in flatten_tree(browser.json)]
        )

        subtasktemplatefolder.getOrdering().moveObjectsToBottom([subtasktemplate.getId()])

        browser.open(
            '{}/@task-template-structure'.format(
                self.tasktemplatefolder.absolute_url()),
            headers=self.api_headers)

        self.assertEqual(
            [
                u'Verfahren Neuanstellung',
                u'Arbeitsplatz einrichten.',
                u'sub-tasktemplate folder 1',
                u'sub-task 1.2',
                u'sub-task 1.1',
            ],
            [item.get('title') for item in flatten_tree(browser.json)]
        )

    @browsing
    def test_regular_users_can_fetch_task_template_structures(self, browser):
        self.activate_feature('tasktemplatefolder_nesting')
        self.login(self.regular_user, browser=browser)

        browser.open(
            '{}/@task-template-structure'.format(
                self.tasktemplatefolder.absolute_url()),
            headers=self.api_headers)

        self.assertEqual(200, browser.status_code)

    @browsing
    def test_deadline_of_tasktemplate_is_absolute(self, browser):
        self.login(self.administrator, browser=browser)

        self.tasktemplate.deadline = 5

        with freeze(datetime(2021, 12, 10)):  # Friday
            browser.open(
                '{}/@task-template-structure'.format(
                    self.tasktemplatefolder.absolute_url()),
                headers=self.api_headers)

        self.assertEqual(u'2021-12-17', browser.json.get('items')[0].get('deadline'))

    @browsing
    def test_include_static_is_private_attribute_for_tasktemplates(self, browser):
        self.login(self.administrator, browser=browser)

        browser.open(
            '{}/@task-template-structure'.format(
                self.tasktemplatefolder.absolute_url()),
            headers=self.api_headers)

        self.assertEqual(False, browser.json.get('items')[0].get('is_private'))
