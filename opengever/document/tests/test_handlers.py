from datetime import datetime
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docxcompose.properties import CustomProperties
from ftw.builder import Builder
from ftw.builder import create
from ftw.testbrowser import browsing
from opengever.document.interfaces import ICheckinCheckoutManager
from opengever.dossier.docprops import DocPropertyWriter
from opengever.dossier.docprops import TemporaryDocFile
from opengever.journal.handlers import DOC_PROPERTIES_UPDATED
from opengever.journal.tests.utils import get_journal_entry
from opengever.testing import FunctionalTestCase
from plone import api
from plone.app.testing import TEST_USER_ID
from zope.component import getMultiAdapter


class TestHandlers(FunctionalTestCase):

    def setUp(self):
        super(TestHandlers, self).setUp()
        self.setup_fullname(fullname='Peter')
        self.set_docproperty_export_enabled(True)

        self.dossier = create(Builder('dossier').titled(u'Dossier'))
        self.target_dossier = create(Builder('dossier').titled(u'Target'))

        self.doc_with_gever_properties = create(
            Builder('document')
            .within(self.dossier)
            .titled("Document with file")
            .having(document_date=datetime(2010, 12, 30, 0, 0))
            .with_asset_file('with_gever_properties.docx'))

    def create_invalid_docx(self):
        """Produce a file with *.docx extension and mimetype, but contents
        that aren't a valid DOCX, in order to trip up DocProperties update.
        """
        invalid_docx = create(
            Builder('document')
            .within(self.dossier)
            .titled("Invalid DOCX")
            .attach_file_containing('foo', u'invalid.dat'))

        # Only change file type to .docx after creation, because otherwise
        # creation already fails because of the attempt to update docprops
        invalid_docx.file.filename = u'invalid.docx'
        invalid_docx.file.contentType = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # noqa
        return invalid_docx

    def tearDown(self):
        super(TestHandlers, self).tearDown()
        self.set_docproperty_export_enabled(False)

    def assert_doc_properties_updated_journal_entry_generated(self, document, entry=-1):
        entry = get_journal_entry(document, entry=entry)

        self.assertEqual(DOC_PROPERTIES_UPDATED, entry['action']['type'])
        self.assertEqual(TEST_USER_ID, entry['actor'])
        self.assertEqual('', entry['comments'])

    def set_document_property_referencenumber(self):
        DocPropertyWriter(self.doc_with_gever_properties).write_properties(
            False, {'ogg.dossier.reference_number': 'ClientXY / 42'})

    def test_document_checkout_updates_doc_properties(self):
        self.set_document_property_referencenumber()

        manager = getMultiAdapter(
            (self.doc_with_gever_properties, self.request),
            ICheckinCheckoutManager)
        manager.checkout()

        expected_doc_properties = [
            ('Document.ReferenceNumber', 'Client1 / 1 / 1'),
            ('Document.SequenceNumber', '1'),
            ('Dossier.ReferenceNumber', 'Client1 / 1'),
            ('Dossier.Title', 'Dossier'),
            ('ogg.document.document_date', datetime(2010, 12, 30, 0, 0)),
            ('ogg.document.reference_number', 'Client1 / 1 / 1'),
            ('ogg.document.sequence_number', '1'),
            ('ogg.document.title', 'Document with file'),
            ('ogg.document.version_number', 0),
            ('ogg.dossier.reference_number', 'Client1 / 1'),
            ('ogg.dossier.sequence_number', '1'),
            ('ogg.dossier.title', 'Dossier'),
            ('ogg.user.email', 'test@example.org'),
            ('ogg.user.firstname', 'User'),
            ('ogg.user.lastname', 'Test'),
            ('ogg.user.title', 'Test User'),
            ('ogg.user.userid', TEST_USER_ID),
            ('User.FullName', 'Test User'),
            ('User.ID', TEST_USER_ID),
        ]

        with TemporaryDocFile(self.doc_with_gever_properties.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(expected_doc_properties, properties)
        self.assert_doc_properties_updated_journal_entry_generated(
            self.doc_with_gever_properties)

    def test_document_checkin_updates_doc_properties(self):
        manager = getMultiAdapter(
            (self.doc_with_gever_properties, self.request),
            ICheckinCheckoutManager)
        manager.checkout()
        self.set_document_property_referencenumber()
        manager.checkin()

        expected_doc_properties = [
            ('Document.ReferenceNumber', 'Client1 / 1 / 1'),
            ('Document.SequenceNumber', '1'),
            ('Dossier.ReferenceNumber', 'Client1 / 1'),
            ('Dossier.Title', 'Dossier'),
            ('ogg.document.document_date', datetime(2010, 12, 30, 0, 0)),
            ('ogg.document.reference_number', 'Client1 / 1 / 1'),
            ('ogg.document.sequence_number', '1'),
            ('ogg.document.title', 'Document with file'),
            ('ogg.document.version_number', 0),
            ('ogg.dossier.reference_number', 'Client1 / 1'),
            ('ogg.dossier.sequence_number', '1'),
            ('ogg.dossier.title', 'Dossier'),
            ('ogg.user.email', 'test@example.org'),
            ('ogg.user.firstname', 'User'),
            ('ogg.user.lastname', 'Test'),
            ('ogg.user.title', 'Test User'),
            ('ogg.user.userid', TEST_USER_ID),
            ('User.FullName', 'Test User'),
            ('User.ID', TEST_USER_ID),
        ]

        with TemporaryDocFile(self.doc_with_gever_properties.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(expected_doc_properties, properties)

        self.assert_doc_properties_updated_journal_entry_generated(
            self.doc_with_gever_properties, entry=-2)

    @browsing
    def test_copying_documents_updates_doc_properties(self, browser):
        browser.login()
        browser.open(self.doc_with_gever_properties, view='copy_item')
        browser.open(self.target_dossier, view='tabbed_view')
        browser.css('#contentActionMenus a#paste').first.click()

        copied_doc = self.target_dossier.getFirstChild()

        expected_doc_properties = [
            ('Document.ReferenceNumber', 'Client1 / 2 / 2'),
            ('Document.SequenceNumber', '2'),
            ('Dossier.ReferenceNumber', 'Client1 / 2'),
            ('Dossier.Title', 'Target'),
            ('ogg.document.document_date', datetime(2010, 12, 30, 0, 0)),
            ('ogg.document.reference_number', 'Client1 / 2 / 2'),
            ('ogg.document.sequence_number', '2'),
            ('ogg.document.title', 'copy of Document with file'),
            ('ogg.document.version_number', 0),
            ('ogg.dossier.reference_number', 'Client1 / 2'),
            ('ogg.dossier.sequence_number', '2'),
            ('ogg.dossier.title', 'Target'),
            ('ogg.user.email', 'test@example.org'),
            ('ogg.user.firstname', 'User'),
            ('ogg.user.lastname', 'Test'),
            ('ogg.user.title', 'Test User'),
            ('ogg.user.userid', TEST_USER_ID),
            ('User.FullName', 'Test User'),
            ('User.ID', TEST_USER_ID),
        ]

        with TemporaryDocFile(copied_doc.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(expected_doc_properties, properties)
        self.assert_doc_properties_updated_journal_entry_generated(copied_doc)

    @browsing
    def test_copying_dossier_updates_doc_properties_of_contained_document(self, browser):
        browser.login()
        browser.open(self.dossier, view='copy_item')
        browser.open(self.target_dossier, view='tabbed_view')
        browser.css('#contentActionMenus a#paste').first.click()

        copied_dossier = self.target_dossier.getFirstChild()
        copied_doc = copied_dossier.getFirstChild()

        expected_doc_properties = [
            ('Document.ReferenceNumber', 'Client1 / 2.1 / 2'),
            ('Document.SequenceNumber', '2'),
            ('Dossier.ReferenceNumber', 'Client1 / 2.1'),
            ('Dossier.Title', 'Dossier'),
            ('ogg.document.document_date', datetime(2010, 12, 30, 0, 0)),
            ('ogg.document.reference_number', 'Client1 / 2.1 / 2'),
            ('ogg.document.sequence_number', '2'),
            ('ogg.document.title', 'copy of Document with file'),
            ('ogg.document.version_number', 0),
            ('ogg.dossier.reference_number', 'Client1 / 2.1'),
            ('ogg.dossier.sequence_number', '3'),
            ('ogg.dossier.title', 'Dossier'),
            ('ogg.user.email', 'test@example.org'),
            ('ogg.user.firstname', 'User'),
            ('ogg.user.lastname', 'Test'),
            ('ogg.user.title', 'Test User'),
            ('ogg.user.userid', TEST_USER_ID),
            ('User.FullName', 'Test User'),
            ('User.ID', TEST_USER_ID),
        ]

        with TemporaryDocFile(copied_doc.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(expected_doc_properties, properties)
        self.assert_doc_properties_updated_journal_entry_generated(copied_doc)

        # XXX we should also test that the document gets properly reindexed in
        # solr as a regression test for https://github.com/4teamwork/opengever.core/pull/5553

    def test_moving_documents_updates_doc_properties(self):
        api.content.move(source=self.doc_with_gever_properties,
                         target=self.target_dossier)
        moved_doc = self.target_dossier.getFirstChild()

        expected_doc_properties = [
            ('Document.ReferenceNumber', 'Client1 / 2 / 1'),
            ('Document.SequenceNumber', '1'),
            ('Dossier.ReferenceNumber', 'Client1 / 2'),
            ('Dossier.Title', 'Target'),
            ('ogg.document.document_date', datetime(2010, 12, 30, 0, 0)),
            ('ogg.document.reference_number', 'Client1 / 2 / 1'),
            ('ogg.document.sequence_number', '1'),
            ('ogg.document.title', 'Document with file'),
            ('ogg.document.version_number', 0),
            ('ogg.dossier.reference_number', 'Client1 / 2'),
            ('ogg.dossier.sequence_number', '2'),
            ('ogg.dossier.title', 'Target'),
            ('ogg.user.email', 'test@example.org'),
            ('ogg.user.firstname', 'User'),
            ('ogg.user.lastname', 'Test'),
            ('ogg.user.title', 'Test User'),
            ('ogg.user.userid', TEST_USER_ID),
            ('User.FullName', 'Test User'),
            ('User.ID', TEST_USER_ID),
        ]
        with TemporaryDocFile(moved_doc.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(expected_doc_properties, properties)
        self.assert_doc_properties_updated_journal_entry_generated(moved_doc)

    def test_failure_to_update_docprops_doesnt_block_checkout(self):
        invalid_docx = self.create_invalid_docx()
        manager = getMultiAdapter(
            (invalid_docx, self.request),
            ICheckinCheckoutManager)
        manager.checkout()
        self.assertTrue(invalid_docx.is_checked_out())

    def test_failure_to_update_docprops_doesnt_block_checkin(self):
        invalid_docx = self.create_invalid_docx()
        manager = getMultiAdapter(
            (invalid_docx, self.request),
            ICheckinCheckoutManager)
        manager.checkout()
        manager.checkin()
        self.assertFalse(invalid_docx.is_checked_out())

    def test_failure_to_update_docprops_doesnt_block_moving(self):
        invalid_docx = self.create_invalid_docx()
        api.content.move(source=invalid_docx,
                         target=self.target_dossier)
        moved_doc = self.target_dossier.getFirstChild()
        self.assertEqual('Invalid DOCX', moved_doc.title)

    def test_failure_to_update_docprops_does_block_creation_of_new_doc(self):
        with self.assertRaises(PackageNotFoundError):
            create(
                Builder('document')
                .within(self.dossier)
                .titled("Invalid DOCX")
                .attach_file_containing('foo', u'invalid.docx'))

    def test_failure_to_update_docprops_does_block_copying(self):
        invalid_docx = self.create_invalid_docx()
        with self.assertRaises(PackageNotFoundError):
            api.content.copy(source=invalid_docx,
                             target=self.target_dossier)
