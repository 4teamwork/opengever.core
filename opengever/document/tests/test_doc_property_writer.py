from datetime import datetime
from docx import Document
from docxcompose.properties import CustomProperties
from docxcompose.sdt import StructuredDocumentTags
from ftw.builder import Builder
from ftw.builder import create
from opengever.document.docprops import DocPropertyWriter
from opengever.document.docprops import TemporaryDocFile
from opengever.dossier.tests import EXPECTED_DOC_PROPERTIES
from opengever.journal.handlers import DOC_PROPERTIES_UPDATED
from opengever.journal.tests.utils import get_journal_entry
from opengever.journal.tests.utils import get_journal_length
from opengever.testing import assets
from opengever.testing import IntegrationTestCase
from plone.namedfile.file import NamedBlobFile
import pytz


class TestDocPropertyWriterExportDisabled(IntegrationTestCase):

    def test_export_is_disabled_when_feature_is_not_enabled(self):
        self.login(self.regular_user)

        self.assertFalse(DocPropertyWriter(self.document).is_export_enabled())


class TestDocPropertyWriter(IntegrationTestCase):

    maxDiff = None

    features = ('doc-properties',)

    def test_export_is_enabled_when_feature_is_enabled(self):
        self.login(self.regular_user)

        self.assertTrue(DocPropertyWriter(self.document).is_export_enabled())

    def test_is_supported_file(self):
        self.login(self.regular_user)

        writer = DocPropertyWriter(self.document)

        self.assertTrue(writer.is_supported_file())

        self.document.file.contentType = 'text/foo'

        self.assertFalse(writer.is_supported_file())

    def test_has_file(self):
        self.login(self.dossier_responsible)
        self.assertTrue(DocPropertyWriter(self.document).has_file())
        self.assertFalse(DocPropertyWriter(self.empty_document).has_file())

    def test_document_with_gever_properties_is_updated_with_all_properties(
            self,
        ):
        self.login(self.regular_user)
        self.with_asset_file('with_gever_properties.docx')

        DocPropertyWriter(self.document).update_doc_properties(
            only_existing=True)

        expected_properties = [
            (key, value) for key, value
            in EXPECTED_DOC_PROPERTIES.items() if value is not None]
        with TemporaryDocFile(self.document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()

            self.assertItemsEqual(expected_properties, properties)

    def test_overwrites_properties_of_wrong_type(self):
        self.login(self.regular_user)
        self.with_asset_file('with_property_of_wrong_type.docx')

        (
            DocPropertyWriter(self.document)
            .update_doc_properties(only_existing=False)
            )

        with TemporaryDocFile(self.document.file) as tmpfile:
            properties = dict(CustomProperties(Document(tmpfile.path)).items())
            self.assertEqual(
                datetime(2010, 1, 3),
                properties['ogg.document.document_date'],
                )

    def test_files_with_custom_properties_are_not_updated(self):
        self.login(self.regular_user)
        self.with_asset_file('with_custom_properties.docx')

        expected_doc_properties = [('Test', 'Peter',)]

        (
            DocPropertyWriter(self.document)
            .update_doc_properties(only_existing=True)
            )

        with TemporaryDocFile(self.document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(expected_doc_properties, properties)

        self.assertEqual(1, get_journal_length(self.document))

        entry = get_journal_entry(self.document)

        self.assertNotEqual(entry['action']['type'], DOC_PROPERTIES_UPDATED)

    def test_files_with_cached_custom_properties_are_updated(self):
        self.login(self.regular_user)

        filename = 'with_cached_custom_properties.docx'
        expected_properties = {
            'Document.ReferenceNumber': '00 / 11 / 222',
            'Document.SequenceNumber': '222',
            'Dossier.ReferenceNumber': '00 / 11',
            'Dossier.Title': 'Foo Dossier',
            'User.FullName': 'Foo Bar',
            'User.ID': 'foo.bar',
            'ogg.document.delivery_date': datetime(2000, 2, 2, 10, 0, tzinfo=pytz.utc),
            'ogg.document.document_author': 'Foo Bar',
            'ogg.document.document_date': datetime(1990, 1, 1, 10, 0, tzinfo=pytz.utc),
            'ogg.document.document_type': 'Foo Docu',
            'ogg.document.reception_date': datetime(2010, 3, 3, 10, 0, tzinfo=pytz.utc),
            'ogg.document.reference_number': '0 / 11 / 222',
            'ogg.document.sequence_number': '222',
            'ogg.document.title': 'cached-docprops',
            'ogg.document.version_number': 333,
            'ogg.dossier.reference_number': '0 / 11',
            'ogg.dossier.sequence_number': '11',
            'ogg.dossier.title': 'Foo Dossier',
            'ogg.user.title': 'foobar',
            'ogg.user.userid': 'foobar',
        }
        properties = dict(CustomProperties(Document(assets.path_to_asset(filename))).items())
        self.assertItemsEqual(expected_properties, properties)

        self.with_asset_file(filename)

        DocPropertyWriter(self.document).update_doc_properties(only_existing=False)

        expected_properties = {
            'Document.ReferenceNumber': 'Client1 1.1 / 1 / 12',
            'Document.SequenceNumber': '12',
            'Dossier.ReferenceNumber': 'Client1 1.1 / 1',
            'Dossier.Title': u'Vertr\xe4ge mit der kantonalen Finanzverwaltung',
            'User.FullName': u'B\xe4rfuss K\xe4thi',
            'User.ID': 'kathi.barfuss',
            'ogg.document.delivery_date': datetime(2010, 1, 3, 0, 0),
            'ogg.document.document_author': 'test_user_1_',
            'ogg.document.document_date': datetime(2010, 1, 3, 0, 0),
            'ogg.document.document_type': 'Contract',
            'ogg.document.reception_date': datetime(2010, 1, 3, 0, 0),
            'ogg.document.reference_number': 'Client1 1.1 / 1 / 12',
            'ogg.document.sequence_number': '12',
            'ogg.document.title': u'Vertr\xe4gsentwurf',
            'ogg.document.version_number': 0,
            'ogg.dossier.external_reference': u'qpr-900-9001-\xf7',
            'ogg.dossier.reference_number': 'Client1 1.1 / 1',
            'ogg.dossier.sequence_number': '1',
            'ogg.dossier.title': u'Vertr\xe4ge mit der kantonalen Finanzverwaltung',
            'ogg.user.address1': 'Kappelenweg 13',
            'ogg.user.address2': 'Postfach 1234',
            'ogg.user.city': 'Vorkappelen',
            'ogg.user.country': 'Schweiz',
            'ogg.user.department': 'Staatskanzlei',
            'ogg.user.department_abbr': 'SK',
            'ogg.user.description': 'nix',
            'ogg.user.directorate': 'Staatsarchiv',
            'ogg.user.directorate_abbr': 'Arch',
            'ogg.user.email': 'foo@example.com',
            'ogg.user.email2': 'bar@example.com',
            'ogg.user.firstname': u'K\xe4thi',
            'ogg.user.lastname': u'B\xe4rfuss',
            'ogg.user.phone_fax': '012 34 56 77',
            'ogg.user.phone_mobile': '012 34 56 76',
            'ogg.user.phone_office': '012 34 56 78',
            'ogg.user.salutation': 'Prof. Dr.',
            'ogg.user.title': u'B\xe4rfuss K\xe4thi',
            'ogg.user.url': 'http://www.example.com',
            'ogg.user.userid': 'kathi.barfuss',
            'ogg.user.zip_code': '1234',
            'ogg.document.creator.user.email': 'robert.ziegler@gever.local',
            'ogg.document.creator.user.firstname': 'Robert',
            'ogg.document.creator.user.lastname': 'Ziegler',
            'ogg.document.creator.user.title': 'Ziegler Robert',
            'ogg.document.creator.user.userid': 'robert.ziegler',
        }

        with TemporaryDocFile(self.document.file) as tmpfile:
            properties = dict(CustomProperties(Document(tmpfile.path)).items())

        self.assertItemsEqual(expected_properties, properties)

    def test_properties_can_be_added_to_file_without_properties(self):
        self.login(self.regular_user)
        self.with_asset_file('without_custom_properties.docx')

        (
            DocPropertyWriter(self.document)
            .update_doc_properties(only_existing=False)
            )

        expected_properties = {
            key: value for key, value
            in EXPECTED_DOC_PROPERTIES.items() if value is not None}

        with TemporaryDocFile(self.document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertItemsEqual(expected_properties.items(), properties)

    def test_writes_additional_recipient_property_providers(self):
        self.login(self.regular_user)
        self.with_asset_file('without_custom_properties.docx')

        peter = create(
            Builder('person')
            .having(
                firstname=u'Peter',
                lastname=u'M\xfcller',
                )
            )

        address = create(
            Builder('address')
            .for_contact(peter)
            .labeled(u'Home')
            .having(
                street=u'Musterstrasse 283',
                zip_code=u'1234',
                city=u'Hinterkappelen',
                country=u'Schweiz',
                )
            )

        writer = DocPropertyWriter(
            self.document,
            recipient_data=(peter, address),
            )

        writer.update_doc_properties(only_existing=False)

        additional_recipient_properties = {
            'ogg.recipient.contact.title': u'M\xfcller Peter',
            'ogg.recipient.person.firstname': 'Peter',
            'ogg.recipient.person.lastname': u'M\xfcller',
            'ogg.recipient.address.street': u'Musterstrasse 283',
            'ogg.recipient.address.zip_code': '1234',
            'ogg.recipient.address.city': 'Hinterkappelen',
            'ogg.recipient.address.country': 'Schweiz',
            }

        with TemporaryDocFile(self.document.file) as tmpfile:
            properties = CustomProperties(Document(tmpfile.path)).items()
            self.assertDictContainsSubset(
                additional_recipient_properties,
                dict(properties),
                )

    def test_document_with_content_controls_gets_updated(self):
        self.login(self.regular_user)
        self.with_asset_file('content_controls.docx')

        prop_writer = DocPropertyWriter(self.document)
        prop_writer.update_doc_properties(True)

        expected_properties = {
            'ogg.document.document_date': '03.01.2010',
            'ogg.document.reference_number': u'Client1 1.1 / 1 / 14',
            'ogg.document.title': u'Vertr\xe4gsentwurf',
            'ogg.document.version_number': '0',
            'ogg.dossier.title': u'Vertr\xe4ge mit der kantonalen Finanzverwaltung',
        }

        content_control_properties = {}
        with TemporaryDocFile(self.document.file) as tmpfile:
            doc = Document(tmpfile.path)
            sdt = StructuredDocumentTags(doc)
            for key, value in expected_properties.items():
                content_control_properties[key] = sdt.get_text(key)
        self.assertEqual(content_control_properties, expected_properties)

    def with_asset_file(self, filename):
        self.document.file = NamedBlobFile(
            assets.load(filename), filename=unicode(filename))
